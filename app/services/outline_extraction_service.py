from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from uuid import UUID

from docx import Document
from docx.oxml.ns import qn

from app.models.thesis import SectionCreate, SectionUpdate, SubsectionBase
from app.repositories.documents_repository import DocumentsRepository
from app.repositories.thesis_repository import ThesisRepository


IGNORED_HEADINGS = {
    "indice",
    "índice",
    "tabla de contenido",
    "tabla de contenidos",
    "referencias",
    "bibliografia",
    "bibliografía",
    "references",
    "bibliography",
}


@dataclass
class ExtractedOutlineSection:
    title: str
    level: int
    order: int
    subtitles: list[str] = field(default_factory=list)


class OutlineExtractionService:
    def __init__(
        self,
        documents_repository: DocumentsRepository | None = None,
        thesis_repository: ThesisRepository | None = None,
    ) -> None:
        self.documents_repository = documents_repository or DocumentsRepository()
        self.thesis_repository = thesis_repository or ThesisRepository()

    def extract_and_create(self, document_id: UUID, replace: bool = False) -> dict:
        context = self.documents_repository.get_editable_document_context(document_id)
        tesis_id = context["tesis_id"]
        extracted, first_heading_para_idx = self.extract_from_path(context["path"])
        self._maybe_extract_cover(context["path"], tesis_id, first_heading_para_idx)

        if replace:
            self.thesis_repository.delete_all_sections(tesis_id)
            existing_by_key: dict = {}
            next_order = 1
        else:
            existing_sections = self.thesis_repository.list_sections(tesis_id)
            existing_by_key = {
                self._section_key(section.title, section.level): section
                for section in existing_sections
            }
            next_order = max((s.order for s in existing_sections), default=0) + 1

        # Maps heading level → DB id of the last section inserted at that level.
        # Used to assign parent_id for deeper headings.
        level_last_id: dict[int, UUID] = {}

        created_count = 0
        updated_count = 0
        skipped_count = 0
        sections: list[dict] = []

        for item in extracted:
            parent_id = level_last_id.get(item.level - 1)
            key = self._section_key(item.title, item.level)
            existing = existing_by_key.get(key)

            if existing is None:
                created = self.thesis_repository.create_section(
                    tesis_id,
                    SectionCreate(
                        title=item.title,
                        parent_id=parent_id,
                        level=item.level,
                        content="",
                        order=next_order,
                        subsections=[],
                    ),
                )
                level_last_id[item.level] = created.id
                # Reset deeper levels when a new ancestor is inserted
                for lvl in [l for l in list(level_last_id) if l > item.level]:
                    del level_last_id[lvl]
                existing_by_key[key] = created
                sections.append(self._summary(created.id, item, "created", order=next_order))
                created_count += 1
                next_order += 1
            else:
                # Section already exists — update level_last_id so children can still
                # resolve this section as their parent on subsequent iterations
                level_last_id[item.level] = existing.id
                for lvl in [l for l in list(level_last_id) if l > item.level]:
                    del level_last_id[lvl]
                skipped_count += 1
                sections.append(
                    self._summary(
                        existing.id,
                        item,
                        "skipped",
                        order=existing.order,
                        reason="already_exists",
                    )
                )

        return {
            "document_id": context["document_id"],
            "tesis_id": tesis_id,
            "extracted_count": len(extracted),
            "created_count": created_count,
            "updated_count": updated_count,
            "skipped_count": skipped_count,
            "sections": sections,
        }

    def extract_from_path(self, path) -> tuple[list[ExtractedOutlineSection], int]:
        """Return (sections, first_heading_para_idx).

        first_heading_para_idx is the index in document.paragraphs of the first
        detected heading (0 means the very first paragraph is a heading, i.e. no
        pre-heading cover content).
        """
        document = Document(path)
        sections: list[ExtractedOutlineSection] = []
        order = 1
        first_heading_para_idx = 0
        first_heading_found = False

        for para_idx, paragraph in enumerate(document.paragraphs):
            text = self._clean_spaces(paragraph.text)
            if not text:
                continue

            level = self._heading_level(paragraph, text)
            if level is None:
                continue

            if not first_heading_found:
                first_heading_para_idx = para_idx
                first_heading_found = True

            title = self._strip_numbering(text)
            if self._is_ignored_heading(title):
                continue

            sections.append(ExtractedOutlineSection(title=title, level=level, order=order))
            order += 1

        return self._deduplicate_sections(sections), first_heading_para_idx

    def _heading_level(self, paragraph, text: str) -> int | None:
        pPr = paragraph._p.find(qn("w:pPr"))
        style = paragraph.style
        style_name = self._clean_spaces(getattr(style, "name", ""))
        is_heading_style = bool(
            re.search(r"\b(?:heading|t[ií]tulo|encabezado)\b", style_name, re.IGNORECASE)
        )

        # Strategy 0b FIRST: w:numPr ilvl for list-formatted heading hierarchy.
        # Documents like DANIEL 2023 use a single "Heading 1" style for all levels
        # and encode depth via ilvl (0→H1, 1→H2, 2→H3). outlineLvl is always 0 for
        # "Heading 1" and would mask the actual depth if checked first.
        if is_heading_style and pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                ilvl_el = numPr.find(qn("w:ilvl"))
                if ilvl_el is not None:
                    ilvl_str = ilvl_el.get(qn("w:val"))
                    if ilvl_str is not None:
                        return min(int(ilvl_str) + 1, 6)

        # Strategy 0a: w:outlineLvl in paragraph's own pPr (direct override)
        if pPr is not None:
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                val_str = ol.get(qn("w:val"))
                if val_str is not None and int(val_str) < 9:
                    return min(int(val_str) + 1, 6)

        # Strategy 0c: w:outlineLvl defined in the style (styles.xml)
        if style is not None and style.element is not None:
            style_pPr = style.element.find(qn("w:pPr"))
            if style_pPr is not None:
                ol = style_pPr.find(qn("w:outlineLvl"))
                if ol is not None:
                    val_str = ol.get(qn("w:val"))
                    if val_str is not None and int(val_str) < 9:
                        return min(int(val_str) + 1, 6)

        # Strategy 1: style name regex (English + Spanish)
        style_match = re.search(
            r"\b(?:heading|t[ií]tulo|encabezado)\s*([1-6])\b",
            style_name,
            re.IGNORECASE,
        )
        if style_match:
            return int(style_match.group(1))

        # Strategy 2: numbered text pattern (fallback for manually typed numbering)
        numbered_match = re.match(r"^\s*(\d+(?:\.\d+){0,5})[.)]?\s+\S", text)
        if numbered_match:
            return min(numbered_match.group(1).count(".") + 1, 6)

        # Strategy 3: Capítulo prefix
        if re.match(r"^cap[ií]tulo\s+([ivxlcdm]+|\d+)\b", text, re.IGNORECASE):
            return 1

        return None

    def _maybe_extract_cover(
        self, doc_path: Path, tesis_id: UUID, first_heading_para_idx: int
    ) -> None:
        """If the document has pre-heading content and no cover is registered, extract it."""
        if first_heading_para_idx == 0:
            return
        try:
            thesis = self.thesis_repository.get(tesis_id)
            existing = (thesis.thesis_metadata or {}).get("cover_docx_storage_path")
            if existing:
                return
            cover_bytes = self._extract_cover_bytes(doc_path, first_heading_para_idx)
            if not cover_bytes:
                return
            cover_path = Path(doc_path).parent / f"cover-{tesis_id}.docx"
            cover_path.write_bytes(cover_bytes)
            self.thesis_repository.update_cover_metadata(tesis_id, str(cover_path))
        except Exception:
            pass  # cover extraction is best-effort; don't block section import

    def _extract_cover_bytes(self, path: Path, first_heading_para_idx: int) -> bytes | None:
        """Return bytes of a .docx containing only the paragraphs before first_heading_para_idx."""
        doc = Document(path)
        all_paras = doc.paragraphs

        if first_heading_para_idx == 0:
            return None

        # Require at least one non-empty pre-heading paragraph
        if not any(all_paras[i].text.strip() for i in range(first_heading_para_idx)):
            return None

        body = doc.element.body

        # Find the body child element that is the Nth w:p (= first_heading_para_idx)
        para_count = 0
        cut_element = None
        for child in list(body):
            tag = child.tag if isinstance(child.tag, str) else ""
            if tag.endswith("}p"):
                if para_count == first_heading_para_idx:
                    cut_element = child
                    break
                para_count += 1

        if cut_element is None:
            return None

        # Remove everything from cut_element onward, except w:sectPr (required by Word)
        removing = False
        to_remove = []
        for child in list(body):
            if child is cut_element:
                removing = True
            if removing:
                tag = child.tag if isinstance(child.tag, str) else ""
                if not tag.endswith("}sectPr"):
                    to_remove.append(child)

        for child in to_remove:
            body.remove(child)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _deduplicate_sections(
        self,
        sections: list[ExtractedOutlineSection],
    ) -> list[ExtractedOutlineSection]:
        seen: set[tuple[str, int]] = set()
        result: list[ExtractedOutlineSection] = []
        for section in sections:
            key = self._section_key(section.title, section.level)
            if key in seen:
                continue
            seen.add(key)
            result.append(section)
        return result

    def _summary(
        self,
        section_id: UUID | None,
        item: ExtractedOutlineSection,
        status: str,
        order: int,
        reason: str | None = None,
    ) -> dict:
        return {
            "id": section_id,
            "title": item.title,
            "level": item.level,
            "order": order,
            "source": "heading",
            "status": status,
            "reason": reason,
        }

    def _section_key(self, title: str, level: int) -> tuple[str, int]:
        return (self._normalize(title), level)

    def _is_ignored_heading(self, text: str) -> bool:
        return self._normalize(text) in {self._normalize(item) for item in IGNORED_HEADINGS}

    def _strip_numbering(self, text: str) -> str:
        clean = re.sub(r"^\s*\d+(?:\.\d+){0,5}[.)]?\s+", "", text).strip()
        without_chapter = re.sub(
            r"^\s*cap[ií]tulo\s+([ivxlcdm]+|\d+)\s*[:.-]\s*",
            "",
            clean,
            flags=re.IGNORECASE,
        ).strip()
        return without_chapter or clean or text

    def _normalize(self, text: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", self._clean_spaces(text).lower()).strip()

    def _clean_spaces(self, text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()
