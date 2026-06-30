from datetime import UTC, datetime
from io import BytesIO
from copy import deepcopy
import logging
import os
from pathlib import Path
import re
import shutil
import socket
import subprocess
from types import SimpleNamespace
from tempfile import NamedTemporaryFile, TemporaryDirectory
import unicodedata
from urllib.parse import quote, urlsplit
from urllib.request import urlopen
from uuid import UUID, uuid4
from zipfile import ZipFile, ZIP_DEFLATED
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer

from app.core.config import get_settings
from app.models.documents import DocumentResponse
from app.models.references import CitationStyle, ReferenceRead, ReferenceType
from app.models.thesis import SectionRead, ThesisRead
from app.services.citation_service import CitationService, UnsupportedCitationStyleError
from app.services.template_service import TemplateService

BIBLIOGRAPHY_NS = "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CUSTOM_XML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/customXml"
RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CUSTOM_XML_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"
CUSTOM_XML_PROPS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps"
BIBLIOGRAPHY_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOCM_MIME = "application/vnd.ms-word.document.macroEnabled.12"
PDF_MIME = "application/pdf"
DOCX_MAIN_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
DOCM_MAIN_CONTENT_TYPE = "application/vnd.ms-word.document.macroEnabled.main+xml"
VBA_PROJECT_CONTENT_TYPE = "application/vnd.ms-office.vbaProject"
VBA_PROJECT_REL = "http://schemas.microsoft.com/office/2006/relationships/vbaProject"
BIBLIOGRAPHY_ITEM = "customXml/item1.xml"
BIBLIOGRAPHY_ITEM_PROPS = "customXml/itemProps1.xml"
BIBLIOGRAPHY_ITEM_RELS = "customXml/_rels/item1.xml.rels"
VBA_PROJECT_ITEM = "word/vbaProject.bin"
DOCUMENT_RELS = "word/_rels/document.xml.rels"
CONTENT_TYPES = "[Content_Types].xml"
FIELD_LOCALE = "10250"
CITATION_TOKEN_PATTERN = re.compile(r"(\{\{cite:([A-Za-z0-9_-]+)\}\}|\[cite:([A-Za-z0-9_-]+)\])")

logger = logging.getLogger(__name__)

ET.register_namespace("w", WORD_NS)
ET.register_namespace("b", BIBLIOGRAPHY_NS)
ET.register_namespace("ds", CUSTOM_XML_NS)


class DocxService:
    def __init__(
        self,
        template_service: TemplateService | None = None,
        citation_service: CitationService | None = None,
    ) -> None:
        settings = get_settings()
        self.generated_dir = settings.generated_dir
        self.vba_project_path = settings.vba_project_path
        self.enable_docm_macro = settings.enable_docm_macro
        self.backend_url = settings.backend_url.rstrip("/")
        self.template_service = template_service or TemplateService()
        self.citation_service = citation_service or CitationService()

    def generate(
        self,
        thesis: ThesisRead,
        sections: list[SectionRead],
        references: list[ReferenceRead],
        style: CitationStyle = CitationStyle.APA7,
    ) -> DocumentResponse:
        template_path = self.template_service.get_template_path(style)
        document = Document(template_path)
        self._apply_word_metadata(document, thesis)
        self._apply_document_styles(document)
        self._add_cover(document, thesis)
        self._add_table_of_contents(document, sections)
        reference_tags = self._reference_tags(references)
        self._add_sections(document, sections, references, reference_tags, style)
        self._add_references(document, references, style)
        self._enable_auto_field_update(document)

        self.generated_dir.mkdir(parents=True, exist_ok=True)
        include_macro = self._should_include_bibliography_macro()
        extension = "docm" if include_macro else "docx"
        mime_type = DOCM_MIME if include_macro else DOCX_MIME
        timestamp = datetime.now(UTC).strftime("%Y%m%d%H%M%S")
        filename = f"{self._filename_stem(thesis)}-{timestamp}.{extension}"
        path = self.generated_dir / filename
        document.save(path)
        try:
            self._prepend_uploaded_cover_docx(path, thesis.thesis_metadata)
        except Exception as cover_error:
            import logging
            logging.getLogger(__name__).warning(
                "Cover page could not be prepended, continuing without it: %s", cover_error
            )
        self._render_table_of_contents_before_export(path)
        if include_macro:
            self._write_macro_enabled_package(path)
        self._write_bibliography_metadata(path, references, style, reference_tags)

        return DocumentResponse(
            filename=filename,
            path=path,
            download_url=f"/documents/{filename}",
            generated_at=datetime.now(UTC),
            format=extension,
            mime_type=mime_type,
        )

    def upload_filename(self, thesis: ThesisRead) -> str:
        return f"{self._filename_stem(thesis)}.docx"

    def _filename_stem(self, thesis: ThesisRead) -> str:
        title = thesis.thesis_metadata.get("title") or thesis.thesis_metadata.get("titulo")
        value = str(title or thesis.id)
        ascii_value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode()
        stem = re.sub(r"[^A-Za-z0-9._-]+", "-", ascii_value).strip(".-_")
        return (stem or "tesis")[:120]

    def _apply_word_metadata(self, document: Document, thesis: ThesisRead) -> None:
        metadata = thesis.thesis_metadata
        properties = document.core_properties
        title = metadata.get("title") or metadata.get("titulo")
        author = metadata.get("author") or metadata.get("autor")

        if title:
            properties.title = str(title)
            properties.subject = str(title)
        if author:
            properties.author = str(author)
            properties.last_modified_by = str(author)
        properties.category = "Tesis"
        properties.keywords = "tesis, APA, bibliografia, indice"
        properties.modified = datetime.now(UTC)

    def _apply_document_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        self._set_style_font(normal, bold=False)
        self._set_style_paragraph_format(normal)

        for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            style = document.styles[style_name]
            self._set_style_font(style, bold=True)
            self._set_style_paragraph_format(style)

    def _set_style_font(self, style, bold: bool) -> None:
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)

    def _set_style_paragraph_format(self, style) -> None:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            return

        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _format_paragraph(self, paragraph) -> None:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _reference_tags(self, references: list[ReferenceRead]) -> dict[UUID, str]:
        return {reference.id: f"Ref_{reference.id.hex[:16]}" for reference in references}

    def _add_cover(self, document: Document, thesis: ThesisRead) -> None:
        metadata = thesis.thesis_metadata
        if self._resolve_cover_docx_url(metadata):
            return

        self._format_paragraph(document.add_paragraph(metadata.get("institution") or ""))
        title = document.add_paragraph()
        title.style = document.styles["Title"]
        title.add_run(metadata.get("title") or "Tesis").bold = True
        self._format_paragraph(title)

        if metadata.get("author"):
            self._format_paragraph(document.add_paragraph(f"Autor: {metadata['author']}"))
        if metadata.get("year"):
            self._format_paragraph(document.add_paragraph(str(metadata["year"])))

        document.add_page_break()

    def _prepend_uploaded_cover_docx(self, document_path: Path, metadata: dict) -> None:
        # Try local filesystem path first (avoids HTTP round-trip to backend)
        storage_path = metadata.get("cover_docx_storage_path")
        if storage_path:
            local = Path(str(storage_path))
            if local.is_file():
                cover_type = self._cover_file_type(metadata, str(local))
                if cover_type == "pdf":
                    cover_document = self._pdf_cover_document(local.read_bytes())
                else:
                    cover_document = Document(local)
                self._prepend_cover_document(document_path, cover_document)
                return

        cover_url = self._resolve_cover_docx_url(metadata)
        if not cover_url:
            return

        try:
            parsed = urlsplit(str(cover_url))
            if parsed.scheme not in {"http", "https"}:
                raise ValueError("cover_docx_url must be an HTTP or HTTPS URL")

            with urlopen(str(cover_url), timeout=10) as response:
                payload = response.read()
            if not payload:
                raise ValueError("cover file is empty")

            with TemporaryDirectory() as temp_dir:
                cover_document = self._cover_document_from_payload(
                    payload,
                    metadata,
                    cover_url,
                    Path(temp_dir),
                )
                self._prepend_cover_document(document_path, cover_document)
        except Exception as error:
            raise RuntimeError("No se pudo integrar la carátula") from error

    def _cover_document_from_payload(
        self,
        payload: bytes,
        metadata: dict,
        cover_url: str,
        temp_dir: Path,
    ) -> Document:
        cover_type = self._cover_file_type(metadata, cover_url)
        if cover_type == "pdf":
            return self._pdf_cover_document(payload)

        cover_path = temp_dir / "cover.docx"
        cover_path.write_bytes(payload)
        return Document(cover_path)

    def _cover_file_type(self, metadata: dict, cover_url: str) -> str:
        mime_type = str(metadata.get("cover_docx_mime_type") or "").lower()
        original_name = str(metadata.get("cover_docx_original_name") or "").lower()
        url_path = urlsplit(str(cover_url)).path.lower()

        if mime_type == PDF_MIME or original_name.endswith(".pdf") or url_path.endswith(".pdf"):
            return "pdf"
        return "docx"

    def _pdf_cover_document(self, payload: bytes) -> Document:
        try:
            import fitz
        except ImportError as error:
            raise RuntimeError("PyMuPDF is required to integrate PDF covers") from error

        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(0)
        section.right_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.left_margin = Pt(0)

        with fitz.open(stream=payload, filetype="pdf") as pdf:
            if pdf.page_count < 1:
                raise ValueError("cover PDF has no pages")

            for page_index, page in enumerate(pdf):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_stream = BytesIO(pixmap.tobytes("png"))

                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                page_width = int(section.page_width)
                page_height = int(section.page_height)
                image_width = page_width
                image_height = int(image_width * (pixmap.height / pixmap.width))
                if image_height > page_height:
                    image_height = page_height
                    image_width = int(image_height * (pixmap.width / pixmap.height))

                paragraph.add_run().add_picture(
                    image_stream,
                    width=image_width,
                    height=image_height,
                )

                if page_index < pdf.page_count - 1:
                    document.add_page_break()

        return document

    def _prepend_cover_document(self, document_path: Path, cover_document: Document) -> None:
        cover_document.add_page_break()
        composer = Composer(cover_document)
        composer.append(Document(document_path))

        with NamedTemporaryFile(
            dir=document_path.parent,
            delete=False,
            suffix=document_path.suffix,
        ) as temp_file:
            merged_path = Path(temp_file.name)

        try:
            composer.save(merged_path)
            shutil.move(str(merged_path), document_path)
        finally:
            merged_path.unlink(missing_ok=True)

    def _resolve_cover_docx_url(self, metadata: dict) -> str | None:
        cover_url = metadata.get("cover_docx_url")
        if cover_url:
            return str(cover_url)

        storage_path = metadata.get("cover_docx_storage_path")
        if not storage_path:
            return None

        safe_path = quote(str(storage_path).strip().lstrip("/"), safe="/")
        return f"{self.backend_url}/storage/{safe_path}"

    def _add_table_of_contents(
        self,
        document: Document,
        sections: list[SectionRead],
    ) -> None:
        title = document.add_paragraph()
        title.add_run("Tabla de contenido").bold = True
        self._format_paragraph(title)
        self._add_table_of_contents_field(document, self._table_of_contents_entries(sections))
        document.add_page_break()

    def _enable_auto_field_update(self, document: Document) -> None:
        settings = document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def _add_field(
        self,
        paragraph,
        instruction: str,
        placeholder: str = "",
        dirty: bool = False,
    ) -> None:
        run = paragraph.add_run()

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        if dirty:
            fld_char_begin.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        if placeholder:
            text = OxmlElement("w:t")
            text.text = placeholder
            run._r.append(text)
        run._r.append(fld_char_end)

    def _add_table_of_contents_field(
        self,
        document: Document,
        entries: list[tuple[int, str]],
    ) -> None:
        if not entries:
            entries = [(1, "Sin secciones registradas.")]

        last_index = len(entries) - 1
        for index, (level, text) in enumerate(entries):
            paragraph = document.add_paragraph()
            self._format_toc_paragraph(paragraph, level)

            if index == 0:
                self._append_field_start(
                    paragraph.add_run(),
                    r'TOC \o "1-3" \h \z \u',
                    dirty=True,
                )

            paragraph.add_run(text)

            if index == last_index:
                self._append_field_end(paragraph.add_run())

    def _append_field_start(
        self,
        run,
        instruction: str,
        dirty: bool = False,
    ) -> None:
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        if dirty:
            fld_char_begin.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)

    def _append_field_end(self, run) -> None:
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    def _format_toc_paragraph(self, paragraph, level: int) -> None:
        self._format_paragraph(paragraph)
        paragraph.paragraph_format.left_indent = Pt(18 * max(level - 1, 0))

    def _table_of_contents_entries(
        self,
        sections: list[SectionRead],
    ) -> list[tuple[int, str]]:
        entries: list[tuple[int, str]] = []
        for section in sorted(sections, key=lambda item: (item.order, item.created_at)):
            section_number = str(section.order)
            title = self._toc_text(f"{section_number}. {section.title}")
            if title:
                entries.append((section.level, title))

            for index, subsection in enumerate(self._section_subsections(section), start=1):
                subsection_title = str(subsection.title).strip()
                subsection_content = str(subsection.content or "").strip()
                if not subsection_title and not subsection_content:
                    continue

                title = self._toc_text(
                    f"{section_number}.{index} {subsection_title or 'Subtítulo'}"
                )
                if title:
                    entries.append((min(section.level + 1, 3), title))
        return entries

    def _toc_text(self, text: str) -> str:
        without_citations = CITATION_TOKEN_PATTERN.sub("", str(text))
        return re.sub(r"\s+", " ", without_citations).strip()

    def _add_native_citation_field(
        self,
        paragraph,
        citation_tag: str,
        placeholder: str,
    ) -> None:
        citation_field = OxmlElement("w:fldSimple")
        citation_field.set(
            qn("w:instr"),
            rf"CITATION {citation_tag} \m {citation_tag}",
        )
        citation_field.set(qn("w:dirty"), "true")

        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = placeholder
        run.append(text)
        citation_field.append(run)
        paragraph._p.append(citation_field)

    def _add_sections(
        self,
        document: Document,
        sections: list[SectionRead],
        references: list[ReferenceRead],
        reference_tags: dict[UUID, str],
        style: CitationStyle,
    ) -> None:
        reference_lookup = self._reference_lookup(references, reference_tags, style)
        for section in sorted(sections, key=lambda item: (item.order, item.created_at)):
            section_number = str(section.order)
            self._add_heading(
                document,
                f"{section_number}. {section.title}",
                section.level,
                reference_lookup,
            )

            for block in section.content.split("\n\n"):
                text = block.strip()
                if text:
                    self._add_content_paragraph(document, text, reference_lookup)

            for index, subsection in enumerate(self._section_subsections(section), start=1):
                subsection_title = str(subsection.title).strip()
                subsection_content = str(subsection.content or "").strip()
                if not subsection_title and not subsection_content:
                    continue

                self._add_heading(
                    document,
                    f"{section_number}.{index} {subsection_title or 'Subtítulo'}",
                    min(section.level + 1, 3),
                    reference_lookup,
                )
                for block in subsection_content.split("\n\n"):
                    text = block.strip()
                    if text:
                        self._add_content_paragraph(document, text, reference_lookup)

    def _section_subsections(self, section: SectionRead) -> list[SimpleNamespace]:
        legacy_subtitle = (section.subtitle or "").strip()
        subsections = list(section.subsections or [])
        if legacy_subtitle and not subsections:
            return [SimpleNamespace(title=legacy_subtitle, content="")]
        normalized = []
        for subsection in subsections:
            if isinstance(subsection, dict):
                normalized.append(
                    SimpleNamespace(
                        title=subsection.get("title", ""),
                        content=subsection.get("content", ""),
                    )
                )
                continue
            normalized.append(
                SimpleNamespace(
                    title=getattr(subsection, "title", ""),
                    content=getattr(subsection, "content", ""),
                )
            )
        return normalized

    def _add_heading(
        self,
        document: Document,
        text: str,
        level: int,
        reference_lookup: dict[str, tuple[str, str]],
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.style = document.styles[f"Heading {level}"]
        self._format_paragraph(paragraph)
        self._add_text_with_citations(paragraph, text, reference_lookup)

    def _reference_lookup(
        self,
        references: list[ReferenceRead],
        reference_tags: dict[UUID, str],
        style: CitationStyle,
    ) -> dict[str, tuple[str, str]]:
        lookup: dict[str, tuple[str, str]] = {}
        for index, reference in enumerate(references, start=1):
            reference_id = reference.id
            value = (
                reference_tags[reference_id],
                self._citation_placeholder(reference, style, index),
            )
            lookup[str(reference_id).lower()] = value
            lookup[reference_id.hex.lower()] = value
            lookup[reference_tags[reference_id].lower()] = value
        return lookup

    def _add_content_paragraph(
        self,
        document: Document,
        text: str,
        reference_lookup: dict[str, tuple[str, str]],
    ) -> None:
        paragraph = document.add_paragraph()
        self._format_paragraph(paragraph)
        self._add_text_with_citations(paragraph, text, reference_lookup)

    def _add_text_with_citations(
        self,
        paragraph,
        text: str,
        reference_lookup: dict[str, tuple[str, str]],
    ) -> None:
        cursor = 0
        for match in CITATION_TOKEN_PATTERN.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor : match.start()])

            token_value = (match.group(2) or match.group(3) or "").lower()
            citation = reference_lookup.get(token_value)
            if citation:
                tag, placeholder = citation
                self._add_native_citation_field(paragraph, tag, placeholder)
            else:
                paragraph.add_run(match.group(0))
            cursor = match.end()

        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    def _citation_placeholder(
        self,
        reference: ReferenceRead,
        style: CitationStyle,
        number: int,
    ) -> str:
        if style == CitationStyle.VANCOUVER:
            return f"({number})"
        if style == CitationStyle.IEEE:
            return f"[{number}]"

        authors = [author.last_name for author in reference.authors if author.last_name]
        if not authors:
            author_text = reference.title
        elif len(authors) == 1:
            author_text = authors[0]
        elif len(authors) == 2:
            author_text = f"{authors[0]} & {authors[1]}"
        else:
            author_text = f"{authors[0]} et al."

        if style == CitationStyle.ISO690:
            author_text = author_text.upper()

        year = reference.year if reference.year is not None else "s. f."
        return f"({author_text}, {year})"

    def _add_references(
        self,
        document: Document,
        references: list[ReferenceRead],
        style: CitationStyle,
    ) -> None:
        if document.paragraphs:
            document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        heading = document.add_heading("Referencias", level=1)
        self._format_paragraph(heading)

        entries = self._formatted_reference_entries(references, style)
        last_index = len(entries) - 1
        for index, entry in enumerate(entries):
            paragraph = document.add_paragraph()
            self._format_reference_paragraph(paragraph)

            if index == 0:
                self._append_field_start(
                    paragraph.add_run(),
                    rf"BIBLIOGRAPHY \l {FIELD_LOCALE}",
                    dirty=True,
                )

            paragraph.add_run(entry)

            if index == last_index:
                self._append_field_end(paragraph.add_run())

    def _formatted_reference_entries(
        self,
        references: list[ReferenceRead],
        style: CitationStyle,
    ) -> list[str]:
        if not references:
            return ["Sin referencias registradas."]

        try:
            return self.citation_service.format_references(references, style)
        except UnsupportedCitationStyleError:
            return self.citation_service.format_references(references, CitationStyle.APA7)

    def _format_reference_paragraph(self, paragraph) -> None:
        self._format_paragraph(paragraph)
        paragraph.paragraph_format.left_indent = Pt(36)
        paragraph.paragraph_format.first_line_indent = Pt(-36)

    def _render_table_of_contents_before_export(self, path: Path) -> None:
        if not self._should_render_toc_with_libreoffice():
            return

        soffice_path = shutil.which("libreoffice") or shutil.which("soffice")
        python_path = self._python_with_uno()
        if not soffice_path or not python_path:
            return

        with TemporaryDirectory(prefix="thesis-toc-") as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            rendered_path = temp_dir / path.name
            profile_dir = temp_dir / "lo-profile"
            shutil.copy2(path, rendered_path)
            port = self._free_local_port()
            process = subprocess.Popen(
                [
                    soffice_path,
                    "--headless",
                    f"--accept=socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext",
                    "--nofirststartwizard",
                    "--norestore",
                    "--nodefault",
                    "--nolockcheck",
                    f"-env:UserInstallation={profile_dir.as_uri()}",
                ],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )

            try:
                result = subprocess.run(
                    [
                        python_path,
                        "-c",
                        self._libreoffice_toc_update_script(),
                        str(rendered_path),
                        str(port),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=45,
                    check=False,
                )
                if result.returncode != 0:
                    logger.warning(
                        "No se pudo renderizar la tabla de contenido con LibreOffice: %s",
                        (result.stderr or result.stdout).strip(),
                    )
                    return

                if not self._merge_rendered_table_of_contents(rendered_path, path):
                    logger.warning(
                        "LibreOffice no devolvió una tabla de contenido renderizada para %s",
                        path,
                    )
            except Exception as error:
                logger.warning(
                    "No se pudo actualizar la tabla de contenido antes de exportar: %s",
                    error,
                )
            finally:
                process.terminate()
                try:
                    process.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    process.kill()

    def _should_render_toc_with_libreoffice(self) -> bool:
        value = os.getenv("DOCX_RENDER_TOC_WITH_LIBREOFFICE", "true").strip().lower()
        return value not in {"0", "false", "no", "off"}

    def _python_with_uno(self) -> str | None:
        candidates = [
            os.getenv("LIBREOFFICE_PYTHON"),
            "/usr/bin/python3",
            shutil.which("python3"),
        ]
        seen: set[str] = set()
        for candidate in candidates:
            if not candidate or candidate in seen:
                continue
            seen.add(candidate)
            if not Path(candidate).exists():
                continue
            result = subprocess.run(
                [candidate, "-c", "import uno"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=5,
                check=False,
            )
            if result.returncode == 0:
                return candidate
        return None

    def _free_local_port(self) -> int:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.bind(("127.0.0.1", 0))
            return int(sock.getsockname()[1])

    def _libreoffice_toc_update_script(self) -> str:
        return r'''
import sys
import time
from pathlib import Path

import uno
from com.sun.star.beans import PropertyValue


def prop(name, value):
    item = PropertyValue()
    item.Name = name
    item.Value = value
    return item


path = Path(sys.argv[1]).resolve()
port = sys.argv[2]
context = uno.getComponentContext()
resolver = context.ServiceManager.createInstanceWithContext(
    "com.sun.star.bridge.UnoUrlResolver",
    context,
)

remote_context = None
for _ in range(40):
    try:
        remote_context = resolver.resolve(
            f"uno:socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext"
        )
        break
    except Exception:
        time.sleep(0.25)

if remote_context is None:
    raise RuntimeError("LibreOffice UNO no inició a tiempo")

desktop = remote_context.ServiceManager.createInstanceWithContext(
    "com.sun.star.frame.Desktop",
    remote_context,
)
document = desktop.loadComponentFromURL(
    uno.systemPathToFileUrl(str(path)),
    "_blank",
    0,
    (
        prop("Hidden", True),
        prop("ReadOnly", False),
        prop("UpdateDocMode", 3),
    ),
)
if document is None:
    raise RuntimeError(f"LibreOffice no pudo abrir {path}")

try:
    indexes = document.getDocumentIndexes()
    for index in range(indexes.getCount()):
        indexes.getByIndex(index).update()
    document.getTextFields().refresh()
    document.store()
finally:
    document.close(True)
'''

    def _merge_rendered_table_of_contents(
        self,
        rendered_path: Path,
        target_path: Path,
    ) -> bool:
        with ZipFile(target_path, "r") as archive:
            target_document_xml = archive.read("word/document.xml")
        with ZipFile(rendered_path, "r") as archive:
            rendered_document_xml = archive.read("word/document.xml")

        target_root = ET.fromstring(target_document_xml)
        rendered_root = ET.fromstring(rendered_document_xml)
        target_body = target_root.find(qn("w:body"))
        rendered_body = rendered_root.find(qn("w:body"))
        if target_body is None or rendered_body is None:
            return False

        target_slice = self._toc_body_slice(target_body)
        rendered_slice = self._toc_body_slice(rendered_body)
        if not target_slice or not rendered_slice:
            return False

        rendered_elements = list(rendered_body)[rendered_slice[0] : rendered_slice[1]]
        rendered_text = self._element_text(*rendered_elements)
        if not rendered_elements or "Right-click to update field." in rendered_text:
            return False

        for child in list(target_body)[target_slice[0] : target_slice[1]]:
            target_body.remove(child)

        insert_at = target_slice[0]
        for child in rendered_elements:
            target_body.insert(insert_at, deepcopy(child))
            insert_at += 1

        self._replace_docx_file(
            target_path,
            {"word/document.xml": ET.tostring(target_root, encoding="UTF-8", xml_declaration=True)},
        )
        return True

    def _toc_body_slice(self, body: ET.Element) -> tuple[int, int] | None:
        children = list(body)
        title_index = None
        for index, child in enumerate(children):
            if self._element_text(child).strip() == "Tabla de contenido":
                title_index = index
                break

        if title_index is None:
            return None

        for index in range(title_index + 1, len(children)):
            if self._element_has_page_break(children[index]):
                return (title_index + 1, index)
        return None

    def _element_text(self, *elements: ET.Element) -> str:
        parts: list[str] = []
        for element in elements:
            for text in element.iter(qn("w:t")):
                if text.text:
                    parts.append(text.text)
        return "".join(parts)

    def _element_has_page_break(self, element: ET.Element) -> bool:
        for br in element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    def _replace_docx_file(self, path: Path, replacements: dict[str, bytes]) -> None:
        with NamedTemporaryFile(dir=path.parent, delete=False, suffix=".docx") as tmp:
            tmp_path = Path(tmp.name)

        try:
            with ZipFile(path, "r") as source, ZipFile(tmp_path, "w", ZIP_DEFLATED) as target:
                for item in source.infolist():
                    if item.filename in replacements:
                        continue
                    target.writestr(item, source.read(item.filename))

                for filename, payload in replacements.items():
                    target.writestr(filename, payload)
            tmp_path.replace(path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()

    def _should_include_bibliography_macro(self) -> bool:
        if not self.enable_docm_macro:
            return False
        if self.vba_project_path.exists() and self.vba_project_path.is_file():
            return True

        logger.warning(
            "No se encontró vbaProject.bin en %s; se generará DOCX sin macro.",
            self.vba_project_path,
        )
        return False

    def _write_macro_enabled_package(self, path: Path) -> None:
        with ZipFile(path, "r") as archive:
            names = set(archive.namelist())
            content_types_xml = archive.read(CONTENT_TYPES) if CONTENT_TYPES in names else None
            document_rels_xml = archive.read(DOCUMENT_RELS) if DOCUMENT_RELS in names else None

        replacements = {
            VBA_PROJECT_ITEM: self.vba_project_path.read_bytes(),
            CONTENT_TYPES: self._ensure_macro_content_types_xml(content_types_xml),
            DOCUMENT_RELS: self._ensure_vba_project_relationship_xml(document_rels_xml),
        }
        self._replace_docx_file(path, replacements)

    def _write_bibliography_metadata(
        self,
        path: Path,
        references: list[ReferenceRead],
        style: CitationStyle,
        reference_tags: dict[UUID, str],
    ) -> None:
        with ZipFile(path, "r") as archive:
            names = set(archive.namelist())
            content_types_xml = archive.read(CONTENT_TYPES) if CONTENT_TYPES in names else None
            document_rels_xml = archive.read(DOCUMENT_RELS) if DOCUMENT_RELS in names else None
            item_rels_xml = archive.read(BIBLIOGRAPHY_ITEM_RELS) if BIBLIOGRAPHY_ITEM_RELS in names else None
            item_props_xml = archive.read(BIBLIOGRAPHY_ITEM_PROPS) if BIBLIOGRAPHY_ITEM_PROPS in names else None

        replacements = {
            BIBLIOGRAPHY_ITEM: self._build_sources_xml(references, style, reference_tags),
            BIBLIOGRAPHY_ITEM_PROPS: item_props_xml or self._build_item_props_xml(),
            BIBLIOGRAPHY_ITEM_RELS: self._ensure_item_rels_xml(item_rels_xml),
            CONTENT_TYPES: self._ensure_content_types_xml(content_types_xml),
            DOCUMENT_RELS: self._ensure_document_rels_xml(document_rels_xml),
        }

        with NamedTemporaryFile(dir=path.parent, delete=False, suffix=".docx") as tmp:
            tmp_path = Path(tmp.name)

        try:
            with ZipFile(path, "r") as source, ZipFile(tmp_path, "w", ZIP_DEFLATED) as target:
                for item in source.infolist():
                    if item.filename in replacements:
                        continue
                    target.writestr(item, source.read(item.filename))

                for filename, payload in replacements.items():
                    target.writestr(filename, payload)
            tmp_path.replace(path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()

    def _build_sources_xml(
        self,
        references: list[ReferenceRead],
        style: CitationStyle,
        reference_tags: dict[UUID, str],
    ) -> bytes:
        sources = ET.Element(
            f"{{{BIBLIOGRAPHY_NS}}}Sources",
            {
                "xmlns": BIBLIOGRAPHY_NS,
                "SelectedStyle": self._selected_word_style_path(style),
                "StyleName": self._selected_word_style_name(style),
                "Version": "6",
            },
        )

        for index, reference in enumerate(references, start=1):
            source = ET.SubElement(sources, f"{{{BIBLIOGRAPHY_NS}}}Source")
            self._append_bibliography_text(source, "Tag", reference_tags[reference.id])
            self._append_bibliography_text(source, "SourceType", self._word_source_type(reference.type))
            self._append_bibliography_text(source, "Guid", f"{{{str(reference.id).upper()}}}")
            self._append_bibliography_text(source, "LCID", "0")
            self._append_bibliography_text(source, "RefOrder", str(index))
            self._append_word_authors(source, reference)
            self._append_bibliography_text(source, "Title", reference.title)
            self._append_bibliography_text(source, "Year", str(reference.year) if reference.year else None)
            self._append_bibliography_text(source, "Publisher", reference.publisher)
            self._append_bibliography_text(source, "JournalName", reference.journal)
            self._append_bibliography_text(source, "URL", reference.url)
            self._append_bibliography_text(source, "DOI", reference.doi)

            if reference.accessed_at:
                self._append_bibliography_text(source, "YearAccessed", str(reference.accessed_at.year))
                self._append_bibliography_text(source, "MonthAccessed", str(reference.accessed_at.month))
                self._append_bibliography_text(source, "DayAccessed", str(reference.accessed_at.day))

        return self._xml_bytes(sources)

    def _selected_word_style_path(self, style: CitationStyle) -> str:
        return {
            CitationStyle.APA7: "/APA.XSL",
            CitationStyle.VANCOUVER: "/Vancouver.XSL",
            CitationStyle.IEEE: "/IEEE.XSL",
            CitationStyle.ISO690: "/ISO690.XSL",
        }[style]

    def _selected_word_style_name(self, style: CitationStyle) -> str:
        return {
            CitationStyle.APA7: "APA",
            CitationStyle.VANCOUVER: "Vancouver",
            CitationStyle.IEEE: "IEEE",
            CitationStyle.ISO690: "ISO 690",
        }[style]

    def _word_source_type(self, reference_type: ReferenceType) -> str:
        return {
            ReferenceType.BOOK: "Book",
            ReferenceType.ARTICLE: "JournalArticle",
            ReferenceType.WEB: "InternetSite",
        }[reference_type]

    def _append_word_authors(self, source: ET.Element, reference: ReferenceRead) -> None:
        author_root = ET.SubElement(source, f"{{{BIBLIOGRAPHY_NS}}}Author")
        author = ET.SubElement(author_root, f"{{{BIBLIOGRAPHY_NS}}}Author")
        name_list = ET.SubElement(author, f"{{{BIBLIOGRAPHY_NS}}}NameList")

        for item in reference.authors:
            person = ET.SubElement(name_list, f"{{{BIBLIOGRAPHY_NS}}}Person")
            self._append_bibliography_text(person, "Last", item.last_name)
            self._append_bibliography_text(person, "First", item.first_name)

    def _append_bibliography_text(self, parent: ET.Element, tag: str, value: object | None) -> None:
        if value is None or value == "":
            return
        element = ET.SubElement(parent, f"{{{BIBLIOGRAPHY_NS}}}{tag}")
        element.text = str(value)

    def _build_item_props_xml(self) -> bytes:
        item = ET.Element(
            f"{{{CUSTOM_XML_NS}}}datastoreItem",
            {f"{{{CUSTOM_XML_NS}}}itemID": f"{{{str(uuid4()).upper()}}}"},
        )
        schema_refs = ET.SubElement(item, f"{{{CUSTOM_XML_NS}}}schemaRefs")
        ET.SubElement(schema_refs, f"{{{CUSTOM_XML_NS}}}schemaRef", {f"{{{CUSTOM_XML_NS}}}uri": BIBLIOGRAPHY_NS})
        return self._xml_bytes(item)

    def _ensure_item_rels_xml(self, xml: bytes | None) -> bytes:
        root = self._relationships_root(xml)
        if not self._has_relationship(root, CUSTOM_XML_PROPS_REL, "itemProps1.xml"):
            ET.SubElement(
                root,
                f"{{{RELATIONSHIPS_NS}}}Relationship",
                {
                    "Id": self._next_relationship_id(root),
                    "Type": CUSTOM_XML_PROPS_REL,
                    "Target": "itemProps1.xml",
                },
            )
        return self._xml_bytes(root)

    def _ensure_document_rels_xml(self, xml: bytes | None) -> bytes:
        root = self._relationships_root(xml)
        if not self._has_relationship(root, CUSTOM_XML_REL, f"../{BIBLIOGRAPHY_ITEM}"):
            ET.SubElement(
                root,
                f"{{{RELATIONSHIPS_NS}}}Relationship",
                {
                    "Id": self._next_relationship_id(root),
                    "Type": CUSTOM_XML_REL,
                    "Target": f"../{BIBLIOGRAPHY_ITEM}",
                },
            )
        return self._xml_bytes(root)

    def _ensure_vba_project_relationship_xml(self, xml: bytes | None) -> bytes:
        root = self._relationships_root(xml)
        if not self._has_relationship(root, VBA_PROJECT_REL, "vbaProject.bin"):
            ET.SubElement(
                root,
                f"{{{RELATIONSHIPS_NS}}}Relationship",
                {
                    "Id": self._next_relationship_id(root),
                    "Type": VBA_PROJECT_REL,
                    "Target": "vbaProject.bin",
                },
            )
        return self._xml_bytes(root)

    def _relationships_root(self, xml: bytes | None) -> ET.Element:
        if xml:
            return ET.fromstring(xml)
        return ET.Element(f"{{{RELATIONSHIPS_NS}}}Relationships")

    def _has_relationship(self, root: ET.Element, relationship_type: str, target: str) -> bool:
        for relationship in root.findall(f"{{{RELATIONSHIPS_NS}}}Relationship"):
            if relationship.get("Type") == relationship_type and relationship.get("Target") == target:
                return True
        return False

    def _next_relationship_id(self, root: ET.Element) -> str:
        used = {relationship.get("Id") for relationship in root.findall(f"{{{RELATIONSHIPS_NS}}}Relationship")}
        index = 1
        while f"rId{index}" in used:
            index += 1
        return f"rId{index}"

    def _ensure_content_types_xml(self, xml: bytes | None) -> bytes:
        root = self._content_types_root(xml)
        self._ensure_content_type_override(
            root,
            f"/{BIBLIOGRAPHY_ITEM_PROPS}",
            BIBLIOGRAPHY_CONTENT_TYPE,
        )

        return self._xml_bytes(root)

    def _ensure_macro_content_types_xml(self, xml: bytes | None) -> bytes:
        root = self._content_types_root(xml)
        self._ensure_content_type_override(
            root,
            "/word/document.xml",
            DOCM_MAIN_CONTENT_TYPE,
        )
        self._ensure_content_type_override(
            root,
            f"/{VBA_PROJECT_ITEM}",
            VBA_PROJECT_CONTENT_TYPE,
        )
        return self._xml_bytes(root)

    def _content_types_root(self, xml: bytes | None) -> ET.Element:
        if xml:
            return ET.fromstring(xml)

        root = ET.Element(f"{{{CONTENT_TYPES_NS}}}Types")
        ET.SubElement(root, f"{{{CONTENT_TYPES_NS}}}Default", {"Extension": "xml", "ContentType": "application/xml"})
        ET.SubElement(
            root,
            f"{{{CONTENT_TYPES_NS}}}Default",
            {
                "Extension": "rels",
                "ContentType": "application/vnd.openxmlformats-package.relationships+xml",
            },
        )
        return root

    def _ensure_content_type_override(
        self,
        root: ET.Element,
        part_name: str,
        content_type: str,
    ) -> None:
        for override in root.findall(f"{{{CONTENT_TYPES_NS}}}Override"):
            if override.get("PartName") == part_name:
                override.set("ContentType", content_type)
                return

        ET.SubElement(
            root,
            f"{{{CONTENT_TYPES_NS}}}Override",
            {
                "PartName": part_name,
                "ContentType": content_type,
            },
        )

    def _xml_bytes(self, root: ET.Element) -> bytes:
        payload = ET.tostring(root, encoding="UTF-8", xml_declaration=False)
        return b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + payload
