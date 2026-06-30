from pathlib import Path
import re
from uuid import UUID

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from psycopg.types.json import Jsonb

from app.core.config import get_settings
from app.core.database import get_connection
from app.models.documents import ParagraphEditMode
from app.models.documents import (
    DocumentPreviewBlock,
    DocumentPreviewResponse,
    DocumentRawDataItem,
    DocumentRawParagraph,
    StructuredReferenceRead,
    StructuredReferenceUpdate,
    StructuredSectionRead,
    StructuredSectionUpdate,
)
from app.repositories.errors import RepositoryNotFoundError

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOCM_MIME = "application/vnd.ms-word.document.macroEnabled.12"
NO_EDITABLE_PROGRESS_MESSAGE = "No hay avance editable Word disponible para este documento"


class DocumentsRepository:
    def __init__(self) -> None:
        self.generated_dir = get_settings().generated_dir
        self.schema = get_settings().db_schema

    @property
    def documents_table(self) -> str:
        return f'"{self.schema}".documentos_tesis'

    @property
    def sections_table(self) -> str:
        return f'"{self.schema}".documentos_tesis_secciones'

    @property
    def references_table(self) -> str:
        return f'"{self.schema}".documentos_tesis_referencias'

    def get_path(self, filename: str) -> Path:
        if not re.fullmatch(r"[A-Za-z0-9_.-]+\.(docx|docm)", filename):
            raise FileNotFoundError(filename)

        path = (self.generated_dir / filename).resolve()
        generated_dir = self.generated_dir.resolve()
        if generated_dir not in path.parents or not path.exists():
            raise FileNotFoundError(filename)

        return path

    def register_generated(self, tesis_id: UUID, path: Path, filename: str) -> UUID:
        query = f"""
            INSERT INTO {self.documents_table}
                (tesis_id, subido_por, nombre_archivo, url_archivo_drive,
                 documento_drive_id, version, estado_revision, comentario_revision,
                 ruta_storage, tipo_mime, tamano_bytes)
            VALUES (
                %s,
                NULL,
                %s,
                NULL,
                NULL,
                COALESCE(
                    (
                        SELECT max(version) + 1
                        FROM {self.documents_table}
                        WHERE tesis_id = %s
                    ),
                    1
                ),
                'pendiente',
                'Generado automaticamente por thesis-doc-generator',
                %s,
                %s,
                %s
            )
            RETURNING id
        """
        resolved_path = path.resolve()
        mime_type = self._word_mime_type(filename)
        with get_connection() as connection:
            row = connection.execute(
                query,
                (
                    tesis_id,
                    filename,
                    tesis_id,
                    str(resolved_path),
                    mime_type,
                    resolved_path.stat().st_size,
                ),
            ).fetchone()

        return row["id"]

    def get_raw_data(self, document_id: UUID) -> dict[str, UUID | str | None]:
        return self.get_document_snapshot(document_id)

    def get_raw_document(self, document_id: UUID) -> dict:
        row = self._get_document_row(document_id)
        path = self._accessible_docx_path(row)
        document = Document(path)
        raw_data = self._raw_data_from_document(document)
        return self._raw_document_payload(row["id"], raw_data, document)

    def update_raw_data(self, document_id: UUID, raw_data: str) -> dict[str, UUID | str | None]:
        query = f"""
            UPDATE {self.documents_table}
            SET raw_data = %s, processing_status = 'manual', processing_error = null, actualizado_en = now()
            WHERE id = %s
            RETURNING id, raw_data
        """
        with get_connection() as connection:
            row = connection.execute(query, (raw_data, document_id)).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Document {document_id} was not found")
        return {"document_id": row["id"], "raw_data": row.get("raw_data")}

    def process_document(self, document_id: UUID) -> dict:
        row = self._get_document_row(document_id)
        path = self._accessible_docx_path(row)
        document = Document(path)
        payload = self._parse_document(document, row["id"], path)
        return self.save_structured_document(document_id, payload)

    def save_structured_document(self, document_id: UUID, payload: dict) -> dict:
        snapshot = self._upsert_structured_payload(document_id, payload)
        return snapshot

    def get_document_snapshot(self, document_id: UUID) -> dict:
        row = self._get_document_row(document_id)
        sections = self.list_sections(document_id)
        references = self.list_references(document_id)
        raw_data_json = row.get("raw_data_json")
        if not isinstance(raw_data_json, dict):
            raw_data_json = self._build_raw_data_json(row, sections, references, None)

        return {
            "document_id": row["id"],
            "tesis_id": row["tesis_id"],
            "title": raw_data_json.get("title") if isinstance(raw_data_json, dict) else None,
            "raw_data": row.get("raw_data"),
            "raw_data_json": raw_data_json,
            "processing_status": row.get("processing_status"),
            "processing_error": row.get("processing_error"),
            "sections": [section.model_dump(mode="json") for section in sections],
            "references": [reference.model_dump(mode="json") for reference in references],
        }

    def get_metadata_harness(self, document_id: UUID) -> dict:
        query = f"""
            SELECT
                d.id,
                d.tesis_id,
                d.nombre_archivo,
                d.processing_status,
                f.uname AS format_uname
            FROM {self.documents_table} d
            LEFT JOIN "{self.schema}".tesis t ON t.id = d.tesis_id
            LEFT JOIN "{self.schema}".doc_thesis_formats f ON f.id = t.doc_thesis_format_id
            WHERE d.id = %s
        """
        with get_connection() as connection:
            row = connection.execute(query, (document_id,)).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Document {document_id} was not found")

        sections = self.list_sections(document_id)
        references = self.list_references(document_id)

        word_count = sum(
            len(str(s.content or "").split()) for s in sections
        )

        harness_sections = [
            {
                "heading": s.heading,
                "level": s.level,
                "content": s.content,
                "order": s.order_index,
            }
            for s in sections
        ]

        harness_references = [
            {
                "title": r.title,
                "authors": r.authors,
                "year": r.year,
                "type": r.type,
                "raw_text": r.raw_text,
            }
            for r in references
        ]

        return {
            "document_id": row["id"],
            "title": row.get("nombre_archivo"),
            "format": row.get("format_uname"),
            "word_count": word_count,
            "processing_status": row.get("processing_status") or "pending",
            "sections": harness_sections,
            "references": harness_references,
        }

    def insert_citation(
        self,
        document_id: UUID,
        reference_id: UUID,
        paragraph_index: int,
        char_offset: int,
    ) -> dict:
        row = self._get_document_row(document_id)
        reference = self._ensure_reference_belongs_to_document(row, reference_id)
        path = self._accessible_docx_path(row)
        document = Document(path)
        paragraph = self._paragraph_at(document, paragraph_index)
        self._validate_offset(paragraph, char_offset)

        reference_tag = f"Ref_{reference_id.hex[:16]}"
        citation_text = self._citation_placeholder(reference, reference_tag)
        before = paragraph.text[:char_offset]
        after = paragraph.text[char_offset:]
        self._clear_paragraph_content(paragraph)
        if before:
            paragraph.add_run(before)
        self._add_native_citation_field(paragraph, reference_tag, citation_text)
        if after:
            paragraph.add_run(after)

        self._enable_auto_field_update(document)
        document.save(path)
        return self._refresh_document_row(row["id"], path)

    def insert_heading(
        self,
        document_id: UUID,
        text: str,
        paragraph_index: int,
        char_offset: int,
        level: int,
        mode: ParagraphEditMode,
    ) -> dict:
        return self._insert_styled_paragraph(
            document_id=document_id,
            text=text,
            paragraph_index=paragraph_index,
            char_offset=char_offset,
            style_name=f"Heading {level}",
            mode=mode,
        )

    def insert_subtitle(
        self,
        document_id: UUID,
        text: str,
        paragraph_index: int,
        char_offset: int,
        level: int,
        mode: ParagraphEditMode,
    ) -> dict:
        return self._insert_styled_paragraph(
            document_id=document_id,
            text=text,
            paragraph_index=paragraph_index,
            char_offset=char_offset,
            style_name=f"Heading {level}",
            mode=mode,
        )

    def extract_raw_data(self, document_id: UUID) -> dict[str, UUID | str | None]:
        row = self._get_document_row(document_id)
        path = self._accessible_docx_path(row)
        document = Document(path)
        raw_data = self._raw_data_from_document(document)
        return self.update_raw_data(document_id, raw_data)

    def list_sections(self, document_id: UUID) -> list[StructuredSectionRead]:
        self._get_document_row(document_id)
        query = f"""
            SELECT id, documento_tesis_id AS document_id, heading, level, content,
                   order_index, parent_section_id, source_paragraphs, manual_override,
                   version, created_at, updated_at, deleted_at
            FROM {self.sections_table}
            WHERE documento_tesis_id = %s AND deleted_at IS NULL
            ORDER BY order_index, created_at
        """
        with get_connection() as connection:
            rows = connection.execute(query, (document_id,)).fetchall()

        return [self._section_from_row(row) for row in rows]

    def get_section(self, document_id: UUID, section_id: UUID) -> StructuredSectionRead:
        row = self._get_section_row(document_id, section_id)
        return self._section_from_row(row)

    def update_section(
        self,
        document_id: UUID,
        section_id: UUID,
        payload: StructuredSectionUpdate,
    ) -> StructuredSectionRead:
        self._get_document_row(document_id)
        existing = self._get_section_row(document_id, section_id)
        data = {
            "heading": existing["heading"],
            "level": existing["level"],
            "content": existing["content"],
            "order_index": existing["order_index"],
            "parent_section_id": existing["parent_section_id"],
            "source_paragraphs": Jsonb(existing.get("source_paragraphs") or []),
            "manual_override": bool(existing.get("manual_override")),
        }
        data.update(payload.model_dump(mode="json", exclude_unset=True))
        if "source_paragraphs" in data and not isinstance(data["source_paragraphs"], Jsonb):
            data["source_paragraphs"] = Jsonb(data["source_paragraphs"] or [])

        query = f"""
            UPDATE {self.sections_table}
            SET heading = %(heading)s,
                level = %(level)s,
                content = %(content)s,
                order_index = %(order_index)s,
                parent_section_id = %(parent_section_id)s,
                source_paragraphs = %(source_paragraphs)s,
                manual_override = %(manual_override)s,
                updated_at = now(),
                version = version + 1
            WHERE id = %(section_id)s AND documento_tesis_id = %(document_id)s AND deleted_at IS NULL
            RETURNING id, documento_tesis_id AS document_id, heading, level, content,
                      order_index, parent_section_id, source_paragraphs, manual_override,
                      version, created_at, updated_at, deleted_at
        """
        with get_connection() as connection:
            row = connection.execute(
                query,
                {
                    **data,
                    "section_id": section_id,
                    "document_id": document_id,
                },
            ).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Section {section_id} was not found")

        self._refresh_document_snapshot(document_id)
        return self._section_from_row(row)

    def list_references(self, document_id: UUID) -> list[StructuredReferenceRead]:
        self._get_document_row(document_id)
        query = f"""
            SELECT id, documento_tesis_id AS document_id, type, authors, year, title,
                   raw_text, style, source, confidence, version, created_at, updated_at,
                   deleted_at
            FROM {self.references_table}
            WHERE documento_tesis_id = %s AND deleted_at IS NULL
            ORDER BY lower(title), created_at
        """
        with get_connection() as connection:
            rows = connection.execute(query, (document_id,)).fetchall()

        return [self._reference_from_row(row) for row in rows]

    def get_reference(self, document_id: UUID, reference_id: UUID) -> StructuredReferenceRead:
        row = self._get_reference_row(document_id, reference_id)
        return self._reference_from_row(row)

    def update_reference(
        self,
        document_id: UUID,
        reference_id: UUID,
        payload: StructuredReferenceUpdate,
    ) -> StructuredReferenceRead:
        self._get_document_row(document_id)
        existing = self._get_reference_row(document_id, reference_id)
        data = {
            "type": existing["type"],
            "authors": Jsonb(existing.get("authors") or []),
            "year": existing.get("year"),
            "title": existing["title"],
            "raw_text": existing.get("raw_text") or "",
            "style": existing.get("style"),
            "source": existing.get("source") or "text",
            "confidence": float(existing.get("confidence") or 0.5),
        }
        data.update(payload.model_dump(mode="json", exclude_unset=True))
        if "authors" in data and not isinstance(data["authors"], Jsonb):
            data["authors"] = Jsonb(data["authors"] or [])

        query = f"""
            UPDATE {self.references_table}
            SET type = %(type)s,
                authors = %(authors)s,
                year = %(year)s,
                title = %(title)s,
                raw_text = %(raw_text)s,
                style = %(style)s,
                source = %(source)s,
                confidence = %(confidence)s,
                updated_at = now(),
                version = version + 1
            WHERE id = %(reference_id)s AND documento_tesis_id = %(document_id)s AND deleted_at IS NULL
            RETURNING id, documento_tesis_id AS document_id, type, authors, year, title,
                      raw_text, style, source, confidence, version, created_at, updated_at,
                      deleted_at
        """
        with get_connection() as connection:
            row = connection.execute(
                query,
                {
                    **data,
                    "reference_id": reference_id,
                    "document_id": document_id,
                },
            ).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Reference {reference_id} was not found")

        self._refresh_document_snapshot(document_id)
        return self._reference_from_row(row)

    def get_preview(self, document_id: UUID) -> DocumentPreviewResponse:
        snapshot = self.get_document_snapshot(document_id)
        return self._preview_from_snapshot(snapshot)

    def get_editable_document_context(self, document_id: UUID) -> dict:
        row = self._get_document_row(document_id)
        path = self._accessible_docx_path(row)
        return {"document_id": row["id"], "tesis_id": row["tesis_id"], "path": path}

    def _upsert_structured_payload(self, document_id: UUID, payload: dict) -> dict:
        sections_payload = list(payload.get("sections") or [])
        references_payload = list(payload.get("references") or [])
        title = payload.get("title")

        with get_connection() as connection:
            document_row = connection.execute(
                f"""
                    UPDATE {self.documents_table}
                    SET raw_data = %s,
                        raw_data_json = %s,
                        processing_status = 'processed',
                        processing_error = NULL,
                        processed_at = now(),
                        actualizado_en = now()
                    WHERE id = %s
                    RETURNING id, tesis_id, raw_data, raw_data_json, processing_status, processing_error
                """,
                (
                    payload.get("raw_data"),
                    Jsonb(payload.get("raw_data_json") or {}),
                    document_id,
                ),
            ).fetchone()

            if not document_row:
                raise RepositoryNotFoundError(f"Document {document_id} was not found")

            # Preserve manually-edited section content across reprocessing.
            # The Word file often has empty section bodies, so blindly
            # re-inserting would wipe content the user typed in the editor.
            existing_sections = connection.execute(
                f"""
                    SELECT heading, content, manual_override
                    FROM {self.sections_table}
                    WHERE documento_tesis_id = %s AND deleted_at IS NULL
                """,
                (document_id,),
            ).fetchall()
            preserved_content: dict[str, dict] = {}
            for existing in existing_sections:
                key = self._normalize_text(existing.get("heading") or "")
                if not key:
                    continue
                content = str(existing.get("content") or "").strip()
                if existing.get("manual_override") or content:
                    preserved_content[key] = {
                        "content": existing.get("content") or "",
                        "manual_override": bool(existing.get("manual_override")),
                    }

            # Preserve already-stored references: reprocessing should add new
            # ones, never destroy references the user already has.
            existing_references = connection.execute(
                f"""
                    SELECT title, year, raw_text
                    FROM {self.references_table}
                    WHERE documento_tesis_id = %s AND deleted_at IS NULL
                """,
                (document_id,),
            ).fetchall()
            existing_reference_keys = {
                self._reference_dedupe_key(
                    ref.get("title"), ref.get("year"), ref.get("raw_text")
                )
                for ref in existing_references
            }

            connection.execute(
                f"""
                    UPDATE {self.sections_table}
                    SET deleted_at = now(), updated_at = now(), version = version + 1
                    WHERE documento_tesis_id = %s AND deleted_at IS NULL
                """,
                (document_id,),
            )

            section_id_map: dict[int, UUID] = {}
            for section in sorted(
                sections_payload,
                key=lambda item: int(item.get("order_index") or 0),
            ):
                content = section.get("content") or ""
                manual_override = bool(section.get("manual_override"))
                preserved = preserved_content.get(
                    self._normalize_text(section.get("heading") or "")
                )
                # Restore preserved content when it was manually edited or when
                # the freshly parsed body is empty (nothing to overwrite with).
                if preserved and (preserved["manual_override"] or not str(content).strip()):
                    content = preserved["content"]
                    manual_override = manual_override or preserved["manual_override"]

                row = connection.execute(
                    f"""
                        INSERT INTO {self.sections_table}
                            (documento_tesis_id, heading, level, content, order_index,
                             parent_section_id, source_paragraphs, manual_override)
                        VALUES (%(document_id)s, %(heading)s, %(level)s, %(content)s,
                                %(order_index)s, %(parent_section_id)s, %(source_paragraphs)s,
                                %(manual_override)s)
                        RETURNING id
                    """,
                    {
                        "document_id": document_id,
                        "heading": section.get("heading") or "",
                        "level": int(section.get("level") or 1),
                        "content": content,
                        "order_index": int(section.get("order_index") or 0),
                        "parent_section_id": section.get("parent_section_id"),
                        "source_paragraphs": Jsonb(section.get("source_paragraphs") or []),
                        "manual_override": manual_override,
                    },
                ).fetchone()
                if row:
                    section_id_map[int(section.get("order_index") or 0)] = row["id"]

            for section in sorted(
                sections_payload,
                key=lambda item: int(item.get("order_index") or 0),
            ):
                parent_order_index = section.get("parent_order_index")
                if parent_order_index is None:
                    continue
                parent_section_id = section_id_map.get(int(parent_order_index))
                section_id = section_id_map.get(int(section.get("order_index") or 0))
                if section_id and parent_section_id:
                    connection.execute(
                        f"""
                            UPDATE {self.sections_table}
                            SET parent_section_id = %s, updated_at = now()
                            WHERE id = %s
                        """,
                        (parent_section_id, section_id),
                    )

            for reference in references_payload:
                key = self._reference_dedupe_key(
                    reference.get("title"),
                    reference.get("year"),
                    reference.get("raw_text"),
                )
                if key in existing_reference_keys:
                    continue
                existing_reference_keys.add(key)
                connection.execute(
                    f"""
                        INSERT INTO {self.references_table}
                            (documento_tesis_id, type, authors, year, title, raw_text,
                             style, source, confidence)
                        VALUES (%(document_id)s, %(type)s, %(authors)s, %(year)s, %(title)s,
                                %(raw_text)s, %(style)s, %(source)s, %(confidence)s)
                    """,
                    {
                        "document_id": document_id,
                        "type": reference.get("type") or "reference",
                        "authors": Jsonb(reference.get("authors") or []),
                        "year": reference.get("year"),
                        "title": reference.get("title") or "",
                        "raw_text": reference.get("raw_text") or "",
                        "style": reference.get("style"),
                        "source": reference.get("source") or "text",
                        "confidence": float(reference.get("confidence") or 0.5),
                    },
                )

        return self.get_document_snapshot(document_id)

    def _refresh_document_snapshot(self, document_id: UUID) -> dict:
        row = self._get_document_row(document_id)
        sections = self.list_sections(document_id)
        references = self.list_references(document_id)
        raw_data_json = self._build_raw_data_json(row, sections, references, None)
        raw_data = self._render_flat_text(raw_data_json)
        payload = {
            "document_id": document_id,
            "title": raw_data_json.get("title"),
            "raw_data": raw_data,
            "raw_data_json": raw_data_json,
            "processing_status": row.get("processing_status") or "processed",
            "processing_error": row.get("processing_error"),
            "sections": [section.model_dump(mode="json") for section in sections],
            "references": [reference.model_dump(mode="json") for reference in references],
        }
        with get_connection() as connection:
            connection.execute(
                f"""
                    UPDATE {self.documents_table}
                    SET raw_data = %s,
                        raw_data_json = %s,
                        processing_status = %s,
                        processing_error = %s,
                        processed_at = now(),
                        actualizado_en = now()
                    WHERE id = %s
                """,
                (
                    payload["raw_data"],
                    Jsonb(raw_data_json),
                    payload["processing_status"],
                    payload["processing_error"],
                    document_id,
                ),
            )
        return payload

    def _build_raw_data_json(
        self,
        row: dict,
        sections: list[StructuredSectionRead],
        references: list[StructuredReferenceRead],
        paragraphs: list[DocumentRawParagraph] | None,
    ) -> dict:
        stored = row.get("raw_data_json") if isinstance(row.get("raw_data_json"), dict) else {}
        return {
            "document_id": str(row["id"]),
            "title": stored.get("title") or row.get("nombre_archivo"),
            "sections": [
                {
                    "id": str(section.id),
                    "heading": section.heading,
                    "level": section.level,
                    "content": section.content,
                    "order_index": section.order_index,
                    "parent_section_id": str(section.parent_section_id) if section.parent_section_id else None,
                    "source_paragraphs": section.source_paragraphs,
                    "manual_override": section.manual_override,
                }
                for section in sections
            ],
            "references": [
                {
                    "id": str(reference.id),
                    "type": reference.type,
                    "authors": reference.authors,
                    "year": reference.year,
                    "title": reference.title,
                    "raw_text": reference.raw_text,
                    "style": reference.style,
                    "source": reference.source,
                    "confidence": reference.confidence,
                }
                for reference in references
            ],
            "paragraphs": [
                paragraph.model_dump(mode="json") for paragraph in (paragraphs or [])
            ],
        }

    def _render_flat_text(self, raw_data_json: dict) -> str:
        parts: list[str] = []
        title = raw_data_json.get("title")
        if title:
            parts.append(str(title))
        for section in raw_data_json.get("sections", []):
            parts.append(f"{'#' * int(section.get('level') or 1)} {section.get('heading')}")
            content = str(section.get("content") or "").strip()
            if content:
                parts.append(content)
        references = raw_data_json.get("references") or []
        if references:
            parts.append("Referencias")
            parts.extend(
                str(reference.get("raw_text") or reference.get("title") or "").strip()
                for reference in references
            )
        return "\n\n".join(part for part in parts if str(part).strip())

    def _preview_from_snapshot(self, snapshot: dict) -> DocumentPreviewResponse:
        blocks: list[DocumentPreviewBlock] = []
        title = snapshot.get("title")
        if title:
            blocks.append(
                DocumentPreviewBlock(kind="title", text=str(title), level=1, order_index=0)
            )
        sections = [
            self._section_from_row(section)
            if not isinstance(section, StructuredSectionRead)
            else section
            for section in snapshot.get("sections", [])
        ]
        references = [
            self._reference_from_row(reference)
            if not isinstance(reference, StructuredReferenceRead)
            else reference
            for reference in snapshot.get("references", [])
        ]
        for section in sections:
            blocks.append(
                DocumentPreviewBlock(
                    kind="heading",
                    text=section.heading,
                    level=section.level,
                    section_id=section.id,
                    order_index=section.order_index,
                )
            )
            if section.content.strip():
                for paragraph in self._split_preview_paragraphs(section.content):
                    blocks.append(
                        DocumentPreviewBlock(
                            kind="paragraph",
                            text=paragraph,
                            section_id=section.id,
                            order_index=section.order_index,
                        )
                    )
        if references:
            blocks.append(
                DocumentPreviewBlock(kind="heading", text="Referencias", level=1)
            )
            for reference in references:
                blocks.append(
                    DocumentPreviewBlock(
                        kind="reference",
                        text=self._format_reference_preview(reference),
                        section_id=None,
                        order_index=None,
                    )
                )

        preview_html = self._render_preview_html(title, sections, references)
        return DocumentPreviewResponse(
            document_id=snapshot["document_id"],
            title=title,
            preview_html=preview_html,
            blocks=blocks,
            sections=sections,
            references=references,
        )

    def _render_preview_html(
        self,
        title: str | None,
        sections: list[StructuredSectionRead],
        references: list[StructuredReferenceRead],
    ) -> str:
        from html import escape

        parts: list[str] = []
        if title:
            parts.append(f"<h1>{escape(str(title))}</h1>")
        for section in sections:
            heading_level = min(section.level + 1, 6)
            parts.append(f"<h{heading_level}>{escape(section.heading)}</h{heading_level}>")
            for paragraph in self._split_preview_paragraphs(section.content):
                parts.append(f"<p>{escape(paragraph)}</p>")
        if references:
            parts.append("<h2>Referencias</h2>")
            parts.append("<ol>")
            for reference in references:
                parts.append(f"<li>{escape(self._format_reference_preview(reference))}</li>")
            parts.append("</ol>")
        return "".join(parts)

    def _split_preview_paragraphs(self, content: str) -> list[str]:
        return [part.strip() for part in re.split(r"\n{2,}", content or "") if part.strip()]

    def _format_reference_preview(self, reference: StructuredReferenceRead) -> str:
        authors = ", ".join(
            str(author.get("last_name") or "").strip()
            for author in reference.authors
            if str(author.get("last_name") or "").strip()
        )
        year = reference.year if reference.year is not None else "s. f."
        return f"{authors or reference.title} ({year}). {reference.title}"

    def _section_from_row(self, row: dict) -> StructuredSectionRead:
        return StructuredSectionRead(
            id=row["id"],
            document_id=row.get("document_id") or row.get("documento_tesis_id"),
            heading=row["heading"],
            level=row["level"],
            content=row.get("content") or "",
            order_index=row.get("order_index") or 0,
            parent_section_id=row.get("parent_section_id"),
            source_paragraphs=row.get("source_paragraphs") or [],
            manual_override=bool(row.get("manual_override")),
            version=row["version"],
            created_at=row["created_at"],
            updated_at=row["updated_at"],
            deleted_at=row.get("deleted_at"),
        )

    def _reference_from_row(self, row: dict) -> StructuredReferenceRead:
        return StructuredReferenceRead(
            id=row["id"],
            document_id=row.get("document_id") or row.get("documento_tesis_id"),
            type=row["type"],
            authors=row.get("authors") or [],
            year=row.get("year"),
            title=row["title"],
            raw_text=row.get("raw_text") or "",
            style=row.get("style"),
            source=row.get("source") or "text",
            confidence=float(row.get("confidence") or 0.5),
            version=row["version"],
            created_at=row["created_at"],
            updated_at=row["updated_at"],
            deleted_at=row.get("deleted_at"),
        )

    def _get_section_row(self, document_id: UUID, section_id: UUID) -> dict:
        query = f"""
            SELECT id, documento_tesis_id, heading, level, content, order_index,
                   parent_section_id, source_paragraphs, manual_override, version,
                   created_at, updated_at, deleted_at
            FROM {self.sections_table}
            WHERE id = %s AND documento_tesis_id = %s AND deleted_at IS NULL
        """
        with get_connection() as connection:
            row = connection.execute(query, (section_id, document_id)).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Section {section_id} was not found")
        return row

    def _get_reference_row(self, document_id: UUID, reference_id: UUID) -> dict:
        query = f"""
            SELECT id, documento_tesis_id, type, authors, year, title, raw_text, style,
                   source, confidence, version, created_at, updated_at, deleted_at
            FROM {self.references_table}
            WHERE id = %s AND documento_tesis_id = %s AND deleted_at IS NULL
        """
        with get_connection() as connection:
            row = connection.execute(query, (reference_id, document_id)).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Reference {reference_id} was not found")
        return row

    def _parse_document(self, document: Document, document_id: UUID, path: Path) -> dict:
        paragraphs: list[DocumentRawParagraph] = []
        sections: list[dict] = []
        references: list[dict] = []
        current_section: dict | None = None
        current_references = False
        preamble: list[str] = []
        title = self._document_title(document)
        reference_candidates: list[dict] = []

        for index, paragraph in enumerate(document.paragraphs):
            text = self._clean_text(paragraph.text)
            if not text:
                continue

            heading_level = self._detect_heading_level(paragraph, text)
            style_name = self._clean_text(getattr(getattr(paragraph, "style", None), "name", "")) or None

            if heading_level and not self._is_ignored_heading(text):
                normalized_heading = self._strip_heading_numbering(text)
                if self._is_reference_heading(normalized_heading):
                    current_references = True
                    current_section = None
                    paragraphs.append(
                        DocumentRawParagraph(
                            paragraph_index=index,
                            text=text,
                            style=style_name,
                            heading_level=heading_level,
                            is_reference=True,
                        )
                    )
                    continue

                current_references = False
                section = {
                    "heading": normalized_heading,
                    "level": heading_level,
                    "content": "",
                    "order_index": len(sections) + 1,
                    "parent_order_index": self._parent_order_index(sections, heading_level),
                    "parent_section_id": None,
                    "source_paragraphs": [index],
                    "manual_override": False,
                }
                sections.append(section)
                current_section = section
                paragraphs.append(
                    DocumentRawParagraph(
                        paragraph_index=index,
                        text=text,
                        style=style_name,
                        heading_level=heading_level,
                        section_order_index=section["order_index"],
                    )
                )
                continue

            if current_references:
                reference_candidates.append(
                    {
                        "text": text,
                        "paragraph_index": index,
                        "style": style_name,
                    }
                )
                paragraphs.append(
                    DocumentRawParagraph(
                        paragraph_index=index,
                        text=text,
                        style=style_name,
                        is_reference=True,
                    )
                )
                continue

            if current_section is None:
                preamble.append(text)
                paragraphs.append(
                    DocumentRawParagraph(
                        paragraph_index=index,
                        text=text,
                        style=style_name,
                    )
                )
                continue

            if current_section["content"]:
                current_section["content"] += "\n\n" + text
            else:
                current_section["content"] = text
            current_section["source_paragraphs"].append(index)
            paragraphs.append(
                DocumentRawParagraph(
                    paragraph_index=index,
                    text=text,
                    style=style_name,
                    heading_level=current_section["level"],
                    section_order_index=current_section["order_index"],
                )
            )

        parsed_references = self._parse_references(reference_candidates, path)
        payload_sections = [
            {
                "heading": section["heading"],
                "level": section["level"],
                "content": section["content"],
                "order_index": section["order_index"],
                "parent_order_index": section["parent_order_index"],
                "parent_section_id": None,
                "source_paragraphs": section["source_paragraphs"],
                "manual_override": section["manual_override"],
            }
            for section in sections
        ]
        raw_data_json = {
            "document_id": str(document_id),
            "title": title,
            "preamble": preamble,
            "sections": payload_sections,
            "references": parsed_references,
            "paragraphs": [paragraph.model_dump(mode="json") for paragraph in paragraphs],
            "metadata": self._document_metadata(document),
        }
        raw_data = self._render_flat_text(raw_data_json)
        raw_data_json["raw_data"] = raw_data
        return {
            "document_id": document_id,
            "title": title,
            "raw_data": raw_data,
            "raw_data_json": raw_data_json,
            "sections": payload_sections,
            "references": parsed_references,
            "paragraphs": paragraphs,
        }

    def _parse_references(self, candidates: list[dict], path: Path) -> list[dict]:
        references = list(self._parse_metadata_references(path))
        for candidate in candidates:
            parsed = self._parse_reference_text(candidate["text"])
            if parsed:
                references.append(
                    {
                        **parsed,
                        "raw_text": candidate["text"],
                        "source": "text",
                        "confidence": 0.75,
                    }
                )
        deduplicated: list[dict] = []
        seen: set[tuple[str, str | None, str | None]] = set()
        for reference in references:
            key = (
                self._normalize_text(reference.get("title") or ""),
                str(reference.get("year")) if reference.get("year") is not None else None,
                self._normalize_text(reference.get("raw_text") or ""),
            )
            if key in seen:
                continue
            seen.add(key)
            deduplicated.append(reference)
        return deduplicated

    def _parse_reference_text(self, text: str) -> dict | None:
        clean = self._clean_text(text)
        if len(clean) < 12:
            return None
        if self._is_reference_placeholder(clean):
            return None

        year_match = re.search(r"\((\d{4})\)", clean)
        year = int(year_match.group(1)) if year_match else None
        authors_text = clean.split("(", 1)[0].strip().rstrip(".")
        title = clean
        if year_match:
            tail = clean[year_match.end() :].strip(" .")
            if tail:
                title = tail.split(".")[0].strip() or tail

        authors = self._parse_authors(authors_text)
        url = self._extract_url(clean)
        doi = self._extract_doi(clean)

        # Require at least one real citation signal so placeholder lines or
        # stray prose (e.g. "Sin referencias registradas.") are not stored as
        # references. APA authors carry initials in first_name.
        has_apa_authors = any(author.get("first_name") for author in authors)
        if not (year or url or doi or has_apa_authors):
            return None

        return {
            "type": "web" if url else "article",
            "authors": authors,
            "year": year,
            "title": title[:500],
            "style": "APA7",
            "source": "text",
            "confidence": 0.65,
        }

    def _is_reference_placeholder(self, text: str) -> bool:
        normalized = self._normalize_text(text)
        placeholders = {
            "sin referencias registradas",
            "sin referencias",
            "no se registraron referencias",
            "no hay referencias",
            "no references",
            "no references registered",
        }
        return normalized in placeholders

    def _parse_metadata_references(self, path: Path) -> list[dict]:
        from zipfile import BadZipFile, ZipFile
        import xml.etree.ElementTree as ET

        try:
            with ZipFile(path) as archive:
                names = [
                    name
                    for name in archive.namelist()
                    if name.startswith("customXml/")
                    and name.endswith(".xml")
                    and "/_rels/" not in name
                    and not name.endswith(".xml.rels")
                    and "itemProps" not in name
                ]
                references: list[dict] = []
                for name in names:
                    try:
                        root = ET.fromstring(archive.read(name))
                    except Exception:
                        continue
                    for source in root.iter():
                        if self._local_name(source.tag) != "Source":
                            continue
                        parsed = self._parse_word_source(source)
                        if parsed:
                            references.append(parsed)
                return references
        except Exception:
            return []

    def _parse_word_source(self, source) -> dict | None:
        title = self._clean_text(self._child_text(source, "Title") or "")
        authors = self._word_authors(source)
        if not title:
            return None

        return {
            "type": self._word_reference_type(source, self._child_text(source, "URL")),
            "authors": authors,
            "year": self._int_or_none(self._child_text(source, "Year")),
            "title": title[:500],
            "raw_text": self._metadata_raw_text(source),
            "style": self._citation_style_from_sources(source),
            "source": "metadata",
            "confidence": 0.9,
        }

    def _document_title(self, document: Document) -> str:
        core_title = self._clean_text(getattr(document.core_properties, "title", "") or "")
        if core_title:
            return core_title
        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)
            if text:
                return text[:200]
        return "Documento"

    def _document_metadata(self, document: Document) -> dict:
        core = document.core_properties
        return {
            "title": self._clean_text(getattr(core, "title", "") or "") or None,
            "author": self._clean_text(getattr(core, "author", "") or "") or None,
            "subject": self._clean_text(getattr(core, "subject", "") or "") or None,
            "keywords": self._clean_text(getattr(core, "keywords", "") or "") or None,
            "created": getattr(core, "created", None).isoformat() if getattr(core, "created", None) else None,
            "modified": getattr(core, "modified", None).isoformat() if getattr(core, "modified", None) else None,
        }

    def _parent_order_index(self, sections: list[dict], level: int) -> int | None:
        for section in reversed(sections):
            if int(section["level"]) < level:
                return int(section["order_index"])
        return None

    def _detect_heading_level(self, paragraph: Paragraph, text: str) -> int | None:
        pPr = paragraph._p.find(qn("w:pPr"))
        style = paragraph.style
        style_name = self._clean_text(getattr(style, "name", ""))
        is_heading_style = bool(
            re.search(r"\b(?:heading|t[ií]tulo|encabezado)\b", style_name, re.IGNORECASE)
        )

        # Strategy 0b FIRST: w:numPr ilvl for list-formatted heading hierarchy.
        # Documents like DANIEL 2023 use a single "Heading 1" style for all levels
        # and encode depth via ilvl (0→H1, 1→H2, 2→H3). outlineLvl is always 0 for
        # "Heading 1" and would mask the actual depth if checked first.
        if is_heading_style and pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                ilvl_el = numPr.find(qn("w:ilvl"))
                if ilvl_el is not None:
                    ilvl_str = ilvl_el.get(qn("w:val"))
                    if ilvl_str is not None:
                        return min(int(ilvl_str) + 1, 6)

        # Strategy 0a: w:outlineLvl in paragraph's own pPr (direct override)
        if pPr is not None:
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                val_str = ol.get(qn("w:val"))
                if val_str is not None and int(val_str) < 9:
                    return min(int(val_str) + 1, 6)

        # Strategy 0c: w:outlineLvl defined in the style (styles.xml)
        if style is not None and style.element is not None:
            style_pPr = style.element.find(qn("w:pPr"))
            if style_pPr is not None:
                ol = style_pPr.find(qn("w:outlineLvl"))
                if ol is not None:
                    val_str = ol.get(qn("w:val"))
                    if val_str is not None and int(val_str) < 9:
                        return min(int(val_str) + 1, 6)

        # Strategy 1: style name (English + Spanish)
        heading_match = re.search(
            r"\b(?:heading|t[ií]tulo|encabezado)\s*([1-6])\b", style_name, re.IGNORECASE
        )
        if heading_match:
            return min(int(heading_match.group(1)), 6)
        if re.search(r"\b(?:heading|t[ií]tulo|encabezado)\b", style_name, re.IGNORECASE):
            return 1

        # Strategy 2: numbered text pattern
        if re.match(r"^\d+(?:\.\d+){0,5}\s+\S", text):
            return min(text.split()[0].count(".") + 1, 6)

        # Strategy 3: Capítulo prefix
        if re.match(r"^cap[ií]tulo\s+([ivxlcdm]+|\d+)\b", text, re.IGNORECASE):
            return 1

        return None

    def _strip_heading_numbering(self, text: str) -> str:
        clean = re.sub(r"^\s*\d+(?:\.\d+){0,5}[.)]?\s+", "", text).strip()
        clean = re.sub(r"^\s*cap[ií]tulo\s+([ivxlcdm]+|\d+)\s*[:.-]?\s*", "", clean, flags=re.IGNORECASE).strip()
        return clean or text

    def _is_reference_heading(self, text: str) -> bool:
        return self._normalize_text(text) in {
            "referencias",
            "bibliografia",
            "bibliografia",
            "references",
            "bibliography",
        }

    def _is_ignored_heading(self, text: str) -> bool:
        return self._normalize_text(text) in {
            "indice",
            "indice",
            "tabla de contenido",
            "tabla de contenidos",
            "referencias",
            "bibliografia",
            "bibliografía",
            "references",
            "bibliography",
        }

    def _clean_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", self._clean_text(text).lower()).strip()

    def _reference_dedupe_key(
        self, title: str | None, year, raw_text: str | None
    ) -> tuple[str, str | None, str]:
        return (
            self._normalize_text(title or ""),
            str(year) if year is not None else None,
            self._normalize_text(raw_text or ""),
        )

    def _local_name(self, tag: str) -> str:
        if "}" in tag:
            return tag.rsplit("}", 1)[-1]
        return tag

    def _child_text(self, element, name: str) -> str | None:
        for child in element.iter():
            if self._local_name(child.tag) == name:
                return self._clean_text(child.text or "")
        return None

    def _word_authors(self, source) -> list[dict]:
        authors: list[dict] = []
        for person in source.iter():
            if self._local_name(person.tag) != "Person":
                continue
            last_name = self._clean_text(self._child_text(person, "Last") or "")
            first_name = self._clean_text(
                " ".join(
                    part
                    for part in [
                        self._child_text(person, "First") or "",
                        self._child_text(person, "Middle") or "",
                    ]
                    if part
                )
            )
            if last_name:
                authors.append({"last_name": last_name, "first_name": first_name or None})
        return authors[:8]

    def _word_reference_type(self, source, url: str | None) -> str:
        source_type = self._normalize_text(self._child_text(source, "SourceType") or "")
        if "book" in source_type:
            return "book"
        if url:
            return "web"
        return "article"

    def _metadata_raw_text(self, source) -> str:
        parts = [
            self._child_text(source, "Title"),
            self._child_text(source, "Year"),
            self._child_text(source, "DOI"),
            self._child_text(source, "URL"),
        ]
        return " | ".join(part for part in parts if part)

    def _citation_style_from_sources(self, source) -> str:
        selected = self._child_text(source, "SelectedStyle") or self._child_text(source, "StyleName")
        return self._clean_text(selected or "APA7") or "APA7"

    def _extract_url(self, text: str) -> str | None:
        match = re.search(r"https?://\S+", text)
        return match.group(0) if match else None

    def _extract_doi(self, text: str) -> str | None:
        match = re.search(r"10\.\d{4,9}/\S+", text)
        return match.group(0) if match else None

    def _int_or_none(self, value: str | None) -> int | None:
        try:
            return int(value) if value not in (None, "") else None
        except ValueError:
            return None

    def _parse_authors(self, authors_text: str) -> list[dict]:
        authors_text = self._clean_text(authors_text)
        if not authors_text:
            return []
        if "," in authors_text:
            last_name = self._clean_text(authors_text.split(",", 1)[0])
            remainder = self._clean_text(authors_text.split(",", 1)[1]) if "," in authors_text else ""
            return [{"last_name": last_name, "first_name": remainder or None}]
        if " y " in authors_text.lower():
            return [
                {"last_name": self._clean_text(part), "first_name": None}
                for part in re.split(r"\s+y\s+", authors_text, flags=re.IGNORECASE)
                if self._clean_text(part)
            ]
        return [{"last_name": authors_text, "first_name": None}]

    def _insert_styled_paragraph(
        self,
        document_id: UUID,
        text: str,
        paragraph_index: int,
        char_offset: int,
        style_name: str,
        mode: ParagraphEditMode,
    ) -> dict:
        row = self._get_document_row(document_id)
        path = self._accessible_docx_path(row)
        document = Document(path)
        paragraph = self._paragraph_at(document, paragraph_index)

        if mode == ParagraphEditMode.REPLACE:
            self._replace_paragraph_text(paragraph, text)
            paragraph.style = document.styles[style_name]
        else:
            self._validate_offset(paragraph, char_offset)
            self._insert_paragraph_at_offset(document, paragraph, char_offset, text, style_name)

        self._enable_auto_field_update(document)
        document.save(path)
        return self._refresh_document_row(row["id"], path)

    def _get_document_row(self, document_id: UUID) -> dict:
        query = f"""
            SELECT id, tesis_id, raw_data, raw_data_json, processing_status, processing_error,
                   processed_at, ruta_storage, nombre_archivo, tipo_mime
            FROM {self.documents_table}
            WHERE id = %s
        """
        with get_connection() as connection:
            row = connection.execute(query, (document_id,)).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Document {document_id} was not found")
        return row

    def _accessible_docx_path(self, row: dict) -> Path:
        raw_path = row.get("ruta_storage")
        if not raw_path:
            raise ValueError(NO_EDITABLE_PROGRESS_MESSAGE)

        path = Path(str(raw_path)).expanduser().resolve()
        mime_type = str(row.get("tipo_mime") or "")
        filename = str(row.get("nombre_archivo") or path.name)
        is_docx = (
            path.suffix.lower() == ".docx"
            or path.suffix.lower() == ".docm"
            or filename.lower().endswith(".docx")
            or filename.lower().endswith(".docm")
            or mime_type in {DOCX_MIME, DOCM_MIME}
        )

        if not is_docx or not path.exists() or not path.is_file():
            raise ValueError(NO_EDITABLE_PROGRESS_MESSAGE)

        return path

    def _word_mime_type(self, filename: str) -> str:
        if filename.lower().endswith(".docm"):
            return DOCM_MIME
        return DOCX_MIME

    def _ensure_reference_belongs_to_document(self, row: dict, reference_id: UUID) -> dict | None:
        query = f"""
            SELECT tesis_id, data
            FROM "{self.schema}".tesis_references
            WHERE id = %s AND deleted_at IS NULL
            LIMIT 1
        """
        with get_connection() as connection:
            reference = connection.execute(query, (reference_id,)).fetchone()

        if not reference:
            raise RepositoryNotFoundError(f"Reference {reference_id} was not found")
        if reference["tesis_id"] != row["tesis_id"]:
            raise ValueError("La referencia no pertenece a la tesis del documento")
        return reference

    def _citation_placeholder(self, reference: dict | None, fallback_tag: str) -> str:
        data = reference.get("data") if reference else None
        if not isinstance(data, dict):
            return f"({fallback_tag}, s. f.)"

        authors = data.get("authors") or []
        last_names = [
            str(author.get("last_name") or "").strip()
            for author in authors
            if isinstance(author, dict) and str(author.get("last_name") or "").strip()
        ]
        title = str(data.get("title") or fallback_tag).strip()
        if not last_names:
            author_text = title
        elif len(last_names) == 1:
            author_text = last_names[0]
        elif len(last_names) == 2:
            author_text = f"{last_names[0]} & {last_names[1]}"
        else:
            author_text = f"{last_names[0]} et al."

        year = data.get("year")
        year_text = str(year) if year is not None and year != "" else "s. f."
        return f"({author_text}, {year_text})"

    def _paragraph_at(self, document: Document, paragraph_index: int) -> Paragraph:
        try:
            return document.paragraphs[paragraph_index]
        except IndexError as exc:
            raise ValueError("paragraph_index está fuera del rango del documento") from exc

    def _validate_offset(self, paragraph: Paragraph, char_offset: int) -> None:
        if char_offset > len(paragraph.text):
            raise ValueError("char_offset está fuera del rango del párrafo")

    def _raw_data_from_document(self, document: Document) -> str:
        return "\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

    def _raw_document_payload(self, document_id: UUID, raw_data: str, document: Document) -> dict:
        return {
            "document_id": document_id,
            "raw_data": raw_data,
            "paragraphs": [
                {
                    "paragraph_index": index,
                    "text": paragraph.text,
                    "char_count": len(paragraph.text),
                    "style": paragraph.style.name if paragraph.style else None,
                }
                for index, paragraph in enumerate(document.paragraphs)
            ],
        }

    def _refresh_document_row(self, document_id: UUID, path: Path) -> dict:
        document = Document(path)
        raw_data = self._raw_data_from_document(document)
        query = f"""
            UPDATE {self.documents_table}
            SET raw_data = %s, tamano_bytes = %s, actualizado_en = now()
            WHERE id = %s
            RETURNING id
        """
        with get_connection() as connection:
            row = connection.execute(
                query,
                (raw_data, path.stat().st_size, document_id),
            ).fetchone()

        if not row:
            raise RepositoryNotFoundError(f"Document {document_id} was not found")
        return self._raw_document_payload(row["id"], raw_data, document)

    def _clear_paragraph_content(self, paragraph: Paragraph) -> None:
        paragraph_element = paragraph._p
        for child in list(paragraph_element):
            if child.tag != qn("w:pPr"):
                paragraph_element.remove(child)

    def _replace_paragraph_text(self, paragraph: Paragraph, text: str) -> None:
        self._clear_paragraph_content(paragraph)
        paragraph.add_run(text)

    def _insert_paragraph_at_offset(
        self,
        document: Document,
        paragraph: Paragraph,
        char_offset: int,
        text: str,
        style_name: str,
    ) -> None:
        current_text = paragraph.text
        before = current_text[:char_offset]
        after = current_text[char_offset:]

        if char_offset == 0:
            inserted = self._insert_paragraph_before(paragraph, text)
            inserted.style = document.styles[style_name]
            return

        if char_offset == len(current_text):
            inserted = self._insert_paragraph_after(paragraph, text)
            inserted.style = document.styles[style_name]
            return

        self._replace_paragraph_text(paragraph, before)
        inserted = self._insert_paragraph_after(paragraph, text)
        inserted.style = document.styles[style_name]
        trailing = self._insert_paragraph_after(inserted, after)
        trailing.style = document.styles["Normal"]

    def _insert_paragraph_before(self, paragraph: Paragraph, text: str) -> Paragraph:
        new_element = OxmlElement("w:p")
        paragraph._p.addprevious(new_element)
        inserted = Paragraph(new_element, paragraph._parent)
        inserted.add_run(text)
        return inserted

    def _insert_paragraph_after(self, paragraph: Paragraph, text: str) -> Paragraph:
        new_element = OxmlElement("w:p")
        paragraph._p.addnext(new_element)
        inserted = Paragraph(new_element, paragraph._parent)
        inserted.add_run(text)
        return inserted

    def _add_native_citation_field(
        self,
        paragraph: Paragraph,
        citation_tag: str,
        citation_text: str,
    ) -> None:
        citation_field = OxmlElement("w:fldSimple")
        citation_field.set(
            qn("w:instr"),
            rf"CITATION {citation_tag} \m {citation_tag}",
        )
        citation_field.set(qn("w:dirty"), "true")

        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = citation_text
        run.append(text)
        citation_field.append(run)
        paragraph._p.append(citation_field)

    def _enable_auto_field_update(self, document: Document) -> None:
        settings = document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
