import asyncio
from datetime import UTC, datetime
from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4
from zipfile import ZIP_DEFLATED, ZipFile

import httpx
from docx import Document

from app.api.routes import documents, references, thesis
from app.services import docx_service
from app.services.outline_extraction_service import OutlineExtractionService
from app.services.reference_extraction_service import ReferenceExtractionService
from app.main import app
from app.models.documents import DocumentResponse
from app.models.references import ReferenceDeleteResponse
from app.models.thesis import SectionRead, ThesisRead
from app.repositories.documents_repository import DocumentsRepository
from app.repositories.thesis_repository import ThesisRepository


def run(coro):
    return asyncio.run(coro)


def make_client() -> httpx.AsyncClient:
    transport = httpx.ASGITransport(app=app)
    return httpx.AsyncClient(transport=transport, base_url="http://test")


class FakeUrlResponse:
    def __init__(self, payload: bytes) -> None:
        self.payload = payload

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False

    def read(self):
        return self.payload


class FakeThesisRepository:
    def create(self, payload):
        now = datetime.now(UTC)
        data = payload.model_dump(mode="json")
        metadata = data.pop("metadata")
        metadata.update(data)
        return ThesisRead(id=uuid4(), data={"metadata": metadata}, version=1, created_at=now, updated_at=now)

    def get(self, tesis_id):
        now = datetime.now(UTC)
        return ThesisRead(id=tesis_id, data={"metadata": {"title": "Demo"}}, version=1, created_at=now, updated_at=now)

    def create_section(self, tesis_id, payload):
        now = datetime.now(UTC)
        return SectionRead(
            id=uuid4(),
            tesis_id=tesis_id,
            version=1,
            created_at=now,
            updated_at=now,
            **payload.model_dump(),
        )

    def list_sections(self, tesis_id):
        return []

    def replace_sections(self, tesis_id, payload):
        return [self.create_section(tesis_id, section) for section in payload]

    def update_section(self, tesis_id, section_id, payload):
        now = datetime.now(UTC)
        data = {
            "title": "Capítulo I actualizado",
            "subtitle": None,
            "level": 1,
            "content": "Texto",
            "order": 1,
        }
        data.update(payload.model_dump(exclude_unset=True))
        return SectionRead(
            id=section_id,
            tesis_id=tesis_id,
            version=2,
            created_at=now,
            updated_at=now,
            **data,
        )

    def append_section_text(self, tesis_id, section_id, text):
        now = datetime.now(UTC)
        return SectionRead(
            id=section_id,
            tesis_id=tesis_id,
            title="Capítulo I",
            subtitle=None,
            level=1,
            content=f"Texto inicial.\n\n{text}",
            order=1,
            version=2,
            created_at=now,
            updated_at=now,
        )

    def delete_section(self, tesis_id, section_id):
        return {"id": section_id, "deleted": True}


class FakeReferencesRepository:
    def create(self, tesis_id, payload):
        now = datetime.now(UTC)
        return {
            "id": uuid4(),
            "tesis_id": tesis_id,
            "version": 1,
            "created_at": now,
            "updated_at": now,
            **payload.model_dump(mode="json"),
        }

    def list_by_thesis(self, tesis_id):
        return []

    def get(self, reference_id):
        now = datetime.now(UTC)
        return {
            "id": reference_id,
            "tesis_id": uuid4(),
            "authors": [{"first_name": "Ada", "last_name": "Lovelace"}],
            "year": 1843,
            "title": "Notas",
            "type": "book",
            "publisher": "Demo",
            "version": 1,
            "created_at": now,
            "updated_at": now,
        }

    def update(self, reference_id, payload):
        data = self.get(reference_id)
        data.update(payload.model_dump(mode="json", exclude_unset=True))
        data["version"] = 2
        return data

    def delete(self, reference_id):
        return ReferenceDeleteResponse(id=reference_id, deleted=True)


class FakeBuilder:
    def build_docx(
        self,
        tesis_id,
        upload_to_backend=False,
        backend_url=None,
        authorization=None,
    ):
        return DocumentResponse(
            filename="demo.docx",
            path=Path("app/generated/demo.docx"),
            download_url="/documents/demo.docx",
            generated_at=datetime.now(UTC),
            uploaded=upload_to_backend,
            upload={"backend_url": backend_url, "authorization": authorization}
            if upload_to_backend
            else None,
        )


class FakeDocumentsRepository:
    raw_data = "Texto crudo"

    def get_path(self, filename):
        return Path("app/generated") / filename

    def get_raw_data(self, document_id):
        return {"document_id": document_id, "raw_data": self.raw_data}

    def get_raw_document(self, document_id):
        return {
            "document_id": document_id,
            "raw_data": self.raw_data,
            "paragraphs": [
                {
                    "paragraph_index": 0,
                    "text": self.raw_data,
                    "char_count": len(self.raw_data),
                    "style": "Normal",
                }
            ],
        }

    def update_raw_data(self, document_id, raw_data):
        self.raw_data = raw_data
        return {"document_id": document_id, "raw_data": raw_data}

    def extract_raw_data(self, document_id):
        self.raw_data = "Texto extraído"
        return {"document_id": document_id, "raw_data": self.raw_data}

    def insert_citation(self, document_id, reference_id, paragraph_index, char_offset):
        self.raw_data = f"Referencia {reference_id} en {paragraph_index}:{char_offset}"
        return self.get_raw_document(document_id)

    def insert_heading(self, document_id, text, paragraph_index, char_offset, level, mode):
        self.raw_data = f"{level}:{mode}:{paragraph_index}:{char_offset}:{text}"
        return self.get_raw_document(document_id)

    def insert_subtitle(self, document_id, text, paragraph_index, char_offset, level, mode):
        self.raw_data = f"{level}:{mode}:{paragraph_index}:{char_offset}:{text}"
        return self.get_raw_document(document_id)


class FakeReferenceExtractionService:
    def extract_and_create(self, document_id):
        return {
            "document_id": document_id,
            "tesis_id": uuid4(),
            "extracted_count": 1,
            "created_count": 1,
            "skipped_count": 0,
        }


class FakeOutlineExtractionService:
    def extract_and_create(self, document_id):
        return {
            "document_id": document_id,
            "tesis_id": uuid4(),
            "extracted_count": 1,
            "created_count": 1,
            "updated_count": 0,
            "skipped_count": 0,
            "sections": [
                {
                    "title": "Capítulo I",
                    "level": 1,
                    "order": 1,
                    "source": "heading",
                    "status": "created",
                    "subtitles": ["Antecedentes"],
                }
            ],
        }


def test_thesis_and_document_routes_use_services(monkeypatch) -> None:
    monkeypatch.setattr(thesis, "get_repository", lambda: FakeThesisRepository())
    monkeypatch.setattr(documents, "get_builder", lambda: FakeBuilder())
    tesis_id = uuid4()

    async def scenario():
        async with make_client() as client:
            created_response = await client.post("/theses", json={"title": "Demo"})
            section_response = await client.post(
                f"/theses/{tesis_id}/sections",
                json={
                    "title": "Capítulo I",
                    "level": 1,
                    "content": "Texto",
                    "order": 1,
                    "subtitle": "Antecedentes",
                    "subsections": [
                        {"title": "Antecedentes", "content": "Contenido 1.1"},
                        {"title": "Marco teórico", "content": "Contenido 1.2"},
                    ],
                },
            )
            index_response = await client.put(
                f"/theses/{tesis_id}/sections",
                json=[
                    {"title": "Capítulo I", "level": 1, "content": "Texto", "order": 1},
                    {"title": "Subtítulo", "level": 2, "content": "Texto", "order": 2},
                ],
            )
            generated_response = await client.post(f"/theses/{tesis_id}/documents/docx")
        return (
            created_response.json(),
            section_response.json(),
            index_response.json(),
            generated_response.json(),
        )

    created, section, index, generated = run(scenario())

    assert created["data"]["metadata"]["title"] == "Demo"
    assert section["title"] == "Capítulo I"
    assert section["subsections"] == [
        {"title": "Antecedentes", "content": "Contenido 1.1"},
        {"title": "Marco teórico", "content": "Contenido 1.2"},
    ]
    assert len(index) == 2
    assert generated["download_url"] == "/documents/demo.docx"


def test_append_section_text_route_uses_repository(monkeypatch) -> None:
    monkeypatch.setattr(thesis, "get_repository", lambda: FakeThesisRepository())
    tesis_id = uuid4()
    section_id = uuid4()

    async def scenario():
        async with make_client() as client:
            response = await client.post(
                f"/theses/{tesis_id}/sections/{section_id}/append-text",
                json={"text": "Texto agregado."},
            )
        return response.json()

    section = run(scenario())

    assert section["id"] == str(section_id)
    assert section["content"] == "Texto inicial.\n\nTexto agregado."
    assert section["version"] == 2


def test_generate_docx_route_can_request_backend_upload(monkeypatch) -> None:
    monkeypatch.setattr(documents, "get_builder", lambda: FakeBuilder())
    tesis_id = uuid4()

    async def scenario():
        async with make_client() as client:
            response = await client.post(
                f"/theses/{tesis_id}/documents/docx",
                json={"upload_to_backend": True},
                headers={
                    "Authorization": "Bearer demo-token",
                    "X-Backend-Base-Url": "http://backend.test",
                },
            )
        return response.json()

    generated = run(scenario())

    assert generated["uploaded"] is True
    assert generated["upload"]["backend_url"] == "http://backend.test"
    assert generated["upload"]["authorization"] == "Bearer demo-token"


def test_repository_returns_legacy_subtitle_as_subsection() -> None:
    now = datetime.now(UTC)
    tesis_id = uuid4()
    section = ThesisRepository()._section_from_row(
        {
            "id": uuid4(),
            "tesis_id": tesis_id,
            "data": {
                "title": "Capítulo I",
                "subtitle": "Antecedentes",
                "level": 1,
                "content": "Texto",
                "order": 1,
            },
            "version": 1,
            "created_at": now,
            "updated_at": now,
            "deleted_at": None,
        }
    )

    assert section.subsections[0].title == "Antecedentes"
    assert section.subsections[0].content == ""


def test_raw_data_routes_use_repository(monkeypatch) -> None:
    repository = FakeDocumentsRepository()
    monkeypatch.setattr(documents, "get_repository", lambda: repository)
    document_id = uuid4()

    async def scenario():
        async with make_client() as client:
            updated_response = await client.patch(
                f"/documents/{document_id}/raw-data",
                json={"raw_data": "Texto manual"},
            )
            read_response = await client.get(f"/documents/{document_id}/raw-data")
            extracted_response = await client.post(
                f"/documents/{document_id}/raw-data/extract"
            )
        return updated_response.json(), read_response.json(), extracted_response.json()

    updated, read, extracted = run(scenario())

    assert updated == {"document_id": str(document_id), "raw_data": "Texto manual"}
    assert read == {"document_id": str(document_id), "raw_data": "Texto manual"}
    assert extracted == {"document_id": str(document_id), "raw_data": "Texto extraído"}


def test_raw_document_and_in_place_edit_routes_use_repository(monkeypatch) -> None:
    repository = FakeDocumentsRepository()
    monkeypatch.setattr(documents, "get_repository", lambda: repository)
    document_id = uuid4()
    reference_id = uuid4()

    async def scenario():
        async with make_client() as client:
            raw_response = await client.get(f"/documents/{document_id}/raw-document")
            citation_response = await client.post(
                f"/documents/{document_id}/citations",
                json={
                    "reference_id": str(reference_id),
                    "paragraph_index": 0,
                    "char_offset": 5,
                },
            )
            heading_response = await client.post(
                f"/documents/{document_id}/headings",
                json={
                    "text": "Capítulo I",
                    "paragraph_index": 0,
                    "char_offset": 0,
                    "level": 1,
                },
            )
            subtitle_response = await client.post(
                f"/documents/{document_id}/subtitles",
                json={
                    "text": "Antecedentes",
                    "paragraph_index": 0,
                    "char_offset": 0,
                    "level": 2,
                    "mode": "replace",
                },
            )
        return (
            raw_response.json(),
            citation_response.json(),
            heading_response.json(),
            subtitle_response.json(),
        )

    raw, citation, heading, subtitle = run(scenario())

    assert raw["paragraphs"][0]["style"] == "Normal"
    assert str(reference_id) in citation["raw_data"]
    assert heading["raw_data"] == "1:insert:0:0:Capítulo I"
    assert subtitle["raw_data"] == "2:replace:0:0:Antecedentes"


def test_repository_extracts_raw_data_from_accessible_docx(tmp_path) -> None:
    path = tmp_path / "demo.docx"
    docx = Document()
    docx.add_paragraph("Primer párrafo")
    docx.add_paragraph("")
    docx.add_paragraph("Segundo párrafo")
    docx.save(path)

    document_id = uuid4()
    repository = DocumentsRepository()
    repository._get_document_row = lambda _document_id: {
        "id": document_id,
        "raw_data": None,
        "ruta_storage": str(path),
        "nombre_archivo": "demo.docx",
        "tipo_mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    repository.update_raw_data = lambda _document_id, raw_data: {
        "document_id": document_id,
        "raw_data": raw_data,
    }

    extracted = repository.extract_raw_data(document_id)

    assert extracted == {
        "document_id": document_id,
        "raw_data": "Primer párrafo\nSegundo párrafo",
    }


def test_repository_rejects_raw_data_without_editable_word_progress() -> None:
    document_id = uuid4()
    repository = DocumentsRepository()
    repository._get_document_row = lambda _document_id: {
        "id": document_id,
        "raw_data": None,
        "ruta_storage": None,
        "nombre_archivo": "avance.pdf",
        "tipo_mime": "application/pdf",
    }

    try:
        repository.extract_raw_data(document_id)
    except ValueError as error:
        assert str(error) == "No hay avance editable Word disponible para este documento"
    else:
        raise AssertionError("Expected ValueError")


def test_reference_extraction_service_creates_references_from_docx_section(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "avance.docx"
    docx = Document()
    docx.add_paragraph("Capítulo I")
    docx.add_paragraph("Referencias")
    docx.add_paragraph(
        "Lovelace, A. (1843). Notas sobre la máquina analítica. Revista Demo. https://doi.org/10.1000/demo"
    )
    docx.add_paragraph(
        "Turing, A. (1950). Computing machinery and intelligence. Mind."
    )
    docx.save(path)

    created = []
    documents_repository = SimpleNamespace(
        get_editable_document_context=lambda _document_id: {
            "document_id": document_id,
            "tesis_id": tesis_id,
            "path": path,
        }
    )
    references_repository = SimpleNamespace(
        list_by_thesis=lambda _tesis_id: [],
        create=lambda _tesis_id, payload: created.append(payload),
    )

    result = ReferenceExtractionService(
        documents_repository=documents_repository,
        references_repository=references_repository,
    ).extract_and_create(document_id)

    assert result["extracted_count"] == 2
    assert result["created_count"] == 2
    assert created[0].authors[0].last_name == "Lovelace"
    assert created[0].authors[0].first_name == "A."
    assert created[0].year == 1843
    assert created[0].title == "Notas sobre la máquina analítica"
    assert created[0].doi == "10.1000/demo"
    assert result["references"][0]["source"] == "text"
    assert result["references"][0]["status"] == "created"


def test_reference_extraction_service_creates_references_from_word_metadata(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "metadata.docx"
    docx = Document()
    docx.add_paragraph("Documento sin bibliografía visible")
    docx.save(path)
    _write_word_sources(
        path,
        """
        <b:Sources xmlns:b="http://schemas.openxmlformats.org/officeDocument/2006/bibliography" SelectedStyle="/APA.XSL" StyleName="APA" Version="6">
          <b:Source>
            <b:Tag>Ref_demo</b:Tag>
            <b:SourceType>Book</b:SourceType>
            <b:Author><b:Author><b:NameList><b:Person><b:Last>Hernández</b:Last><b:First>Roberto</b:First></b:Person></b:NameList></b:Author></b:Author>
            <b:Title>Metodología de la investigación</b:Title>
            <b:Year>2014</b:Year>
            <b:Publisher>McGraw-Hill</b:Publisher>
          </b:Source>
        </b:Sources>
        """,
    )

    created = []
    result = ReferenceExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        references_repository=_references_repo(created),
    ).extract_and_create(document_id)

    assert result["extracted_count"] == 1
    assert result["created_count"] == 1
    assert result["references"][0]["source"] == "metadata"
    assert created[0].type == "book"
    assert created[0].publisher == "McGraw-Hill"
    assert created[0].authors[0].last_name == "Hernández"
    assert created[0].authors[0].first_name == "Roberto"


def test_reference_extraction_service_prioritizes_metadata_over_duplicate_text(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "duplicated.docx"
    docx = Document()
    docx.add_paragraph("Referencias")
    docx.add_paragraph("Hernández, R. (2014). Metodología de la investigación. McGraw-Hill.")
    docx.save(path)
    _write_word_sources(
        path,
        """
        <b:Sources xmlns:b="http://schemas.openxmlformats.org/officeDocument/2006/bibliography" Version="6">
          <b:Source>
            <b:SourceType>Book</b:SourceType>
            <b:Author><b:Author><b:NameList><b:Person><b:Last>Hernández</b:Last><b:First>Roberto</b:First></b:Person></b:NameList></b:Author></b:Author>
            <b:Title>Metodología de la investigación</b:Title>
            <b:Year>2014</b:Year>
            <b:Publisher>McGraw-Hill</b:Publisher>
          </b:Source>
        </b:Sources>
        """,
    )

    created = []
    result = ReferenceExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        references_repository=_references_repo(created),
    ).extract_and_create(document_id)

    assert result["extracted_count"] == 1
    assert result["created_count"] == 1
    assert result["references"][0]["source"] == "metadata"
    assert created[0].type == "book"


def test_reference_extraction_service_reports_existing_references_as_skipped(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "existing.docx"
    docx = Document()
    docx.add_paragraph("Referencias")
    docx.add_paragraph("Turing, A. (1950). Computing machinery and intelligence. Mind.")
    docx.save(path)

    result = ReferenceExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        references_repository=_references_repo(
            created=[],
            existing=[
                SimpleNamespace(
                    title="Computing machinery and intelligence",
                    year=1950,
                    doi=None,
                )
            ],
        ),
    ).extract_and_create(document_id)

    assert result["extracted_count"] == 1
    assert result["created_count"] == 0
    assert result["skipped_count"] == 1
    assert result["references"][0]["status"] == "skipped"
    assert result["references"][0]["reason"] == "already_exists"


def test_reference_extraction_service_handles_docx_without_references(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "empty.docx"
    docx = Document()
    docx.add_paragraph("Capítulo I")
    docx.add_paragraph("Texto sin referencias")
    docx.save(path)

    created = []
    result = ReferenceExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        references_repository=_references_repo(created),
    ).extract_and_create(document_id)

    assert result["extracted_count"] == 0
    assert result["created_count"] == 0
    assert result["skipped_count"] == 0
    assert result["references"] == []


def _documents_context(document_id, tesis_id, path):
    return SimpleNamespace(
        get_editable_document_context=lambda _document_id: {
            "document_id": document_id,
            "tesis_id": tesis_id,
            "path": path,
        }
    )


def _references_repo(created, existing=None):
    return SimpleNamespace(
        list_by_thesis=lambda _tesis_id: existing or [],
        create=lambda _tesis_id, payload: created.append(payload) or SimpleNamespace(id=uuid4()),
    )


def _write_word_sources(path: Path, sources_xml: str) -> None:
    with ZipFile(path, "a", ZIP_DEFLATED) as archive:
        archive.writestr("customXml/item99.xml", sources_xml.strip())


def test_reference_extraction_route_uses_service(monkeypatch) -> None:
    monkeypatch.setattr(
        documents,
        "get_reference_extraction_service",
        lambda: FakeReferenceExtractionService(),
    )
    document_id = uuid4()

    async def scenario():
        async with make_client() as client:
            response = await client.post(f"/documents/{document_id}/references/extract")
        return response.json()

    result = run(scenario())

    assert result["document_id"] == str(document_id)
    assert result["extracted_count"] == 1
    assert result["created_count"] == 1


def test_outline_extraction_service_creates_sections_from_docx_headings(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "outline.docx"
    docx = Document()
    docx.add_paragraph("Tabla de contenido", style="Heading 1")
    docx.add_paragraph("1. Capítulo I", style="Heading 1")
    docx.add_paragraph("1.1 Antecedentes", style="Heading 2")
    docx.add_paragraph("1.2 Marco teórico", style="Heading 2")
    docx.add_paragraph("Referencias", style="Heading 1")
    docx.save(path)

    created = []
    thesis_repository = SimpleNamespace(
        list_sections=lambda _tesis_id: [],
        create_section=lambda _tesis_id, payload: created.append(payload) or _section_read(tesis_id, payload),
        update_section=lambda _tesis_id, _section_id, _payload: None,
    )

    result = OutlineExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        thesis_repository=thesis_repository,
    ).extract_and_create(document_id)

    # Heading 2 paragraphs are now detected as nested sections, not subtitles
    assert result["extracted_count"] == 3
    assert result["created_count"] == 3
    assert result["sections"][0]["title"] == "Capítulo I"
    assert result["sections"][1]["title"] == "Antecedentes"
    assert result["sections"][2]["title"] == "Marco teórico"
    assert created[0].title == "Capítulo I"
    assert created[0].level == 1
    assert created[0].parent_id is None
    assert created[1].title == "Antecedentes"
    assert created[1].level == 2
    assert created[2].title == "Marco teórico"
    assert created[2].level == 2


def test_outline_extraction_service_updates_existing_section_subtitles(tmp_path) -> None:
    tesis_id = uuid4()
    document_id = uuid4()
    path = tmp_path / "outline-existing.docx"
    docx = Document()
    docx.add_paragraph("Capítulo I", style="Heading 1")
    docx.add_paragraph("Antecedentes", style="Heading 2")
    docx.add_paragraph("Marco teórico", style="Heading 2")
    docx.save(path)

    existing = _section_read(
        tesis_id,
        SimpleNamespace(
            title="Capítulo I",
            subtitle=None,
            subsections=[{"title": "Antecedentes", "content": ""}],
            level=1,
            content="",
            order=1,
        ),
    )
    created_sections = []
    thesis_repository = SimpleNamespace(
        list_sections=lambda _tesis_id: [existing],
        create_section=lambda _tesis_id, payload: created_sections.append(payload) or _section_read(tesis_id, payload),
        update_section=lambda _tesis_id, _section_id, payload: existing,
    )

    result = OutlineExtractionService(
        documents_repository=_documents_context(document_id, tesis_id, path),
        thesis_repository=thesis_repository,
    ).extract_and_create(document_id)

    # Capítulo I already exists → skipped; Antecedentes + Marco teórico are new children
    assert result["skipped_count"] == 1
    assert result["created_count"] == 2
    assert len(created_sections) == 2
    assert created_sections[0].title == "Antecedentes"
    assert created_sections[0].level == 2
    assert created_sections[1].title == "Marco teórico"
    assert created_sections[1].level == 2


def test_outline_extraction_route_uses_service(monkeypatch) -> None:
    monkeypatch.setattr(
        documents,
        "get_outline_extraction_service",
        lambda: FakeOutlineExtractionService(),
    )
    document_id = uuid4()

    async def scenario():
        async with make_client() as client:
            response = await client.post(f"/documents/{document_id}/outline/extract")
        return response.json()

    result = run(scenario())

    assert result["document_id"] == str(document_id)
    assert result["created_count"] == 1
    assert result["sections"][0]["subtitles"] == ["Antecedentes"]


def _section_read(tesis_id, payload):
    now = datetime.now(UTC)
    return SectionRead(
        id=uuid4(),
        tesis_id=tesis_id,
        title=payload.title,
        subtitle=payload.subtitle,
        subsections=payload.subsections,
        level=payload.level,
        content=payload.content,
        order=payload.order,
        version=1,
        created_at=now,
        updated_at=now,
    )


def test_repository_inserts_citation_in_same_docx(tmp_path) -> None:
    path = tmp_path / "demo.docx"
    docx = Document()
    docx.add_paragraph("Texto con referencia.")
    docx.save(path)

    document_id = uuid4()
    thesis_id = uuid4()
    reference_id = uuid4()
    repository = DocumentsRepository()
    repository._get_document_row = lambda _document_id: {
        "id": document_id,
        "tesis_id": thesis_id,
        "raw_data": None,
        "ruta_storage": str(path),
        "nombre_archivo": "demo.docx",
        "tipo_mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    repository._ensure_reference_belongs_to_document = lambda _row, _reference_id: {
        "tesis_id": thesis_id,
        "data": {
            "authors": [{"first_name": None, "last_name": "CHICAGO"}],
            "year": 1985,
            "title": "DEMO REFERENCIA",
        },
    }
    repository._refresh_document_row = lambda _document_id, edited_path: repository._raw_document_payload(
        document_id,
        repository._raw_data_from_document(Document(edited_path)),
        Document(edited_path),
    )

    edited = repository.insert_citation(
        document_id=document_id,
        reference_id=reference_id,
        paragraph_index=0,
        char_offset=len("Texto"),
    )

    assert path.exists()
    assert edited["document_id"] == document_id
    assert edited["paragraphs"][0]["text"] == "Texto con referencia."
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode()
        settings_xml = archive.read("word/settings.xml").decode()

    tag = f"Ref_{reference_id.hex[:16]}"
    assert f'CITATION {tag} \\m {tag}' in document_xml
    assert "<w:fldSimple" in document_xml
    assert "(CHICAGO, 1985)" in document_xml
    assert "updateFields" in settings_xml


def test_repository_inserts_and_replaces_headings_in_same_docx(tmp_path) -> None:
    path = tmp_path / "demo.docx"
    docx = Document()
    docx.add_paragraph("Antes despues")
    docx.add_paragraph("Reemplazar")
    docx.save(path)

    document_id = uuid4()
    repository = DocumentsRepository()
    repository._get_document_row = lambda _document_id: {
        "id": document_id,
        "tesis_id": uuid4(),
        "raw_data": None,
        "ruta_storage": str(path),
        "nombre_archivo": "demo.docx",
        "tipo_mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    repository._refresh_document_row = lambda _document_id, edited_path: repository._raw_document_payload(
        document_id,
        repository._raw_data_from_document(Document(edited_path)),
        Document(edited_path),
    )

    repository.insert_heading(
        document_id=document_id,
        text="Capítulo I",
        paragraph_index=0,
        char_offset=len("Antes"),
        level=1,
        mode="insert",
    )
    edited = repository.insert_subtitle(
        document_id=document_id,
        text="Antecedentes",
        paragraph_index=3,
        char_offset=0,
        level=2,
        mode="replace",
    )

    document = Document(path)
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Antes",
        "Capítulo I",
        " despues",
        "Antecedentes",
    ]
    assert document.paragraphs[1].style.name == "Heading 1"
    assert document.paragraphs[3].style.name == "Heading 2"
    assert edited["document_id"] == document_id


def test_docx_service_prepends_uploaded_cover_docx(tmp_path, monkeypatch) -> None:
    cover_path = tmp_path / "cover.docx"
    cover_docx = Document()
    cover_docx.add_paragraph("CARATULA SUBIDA")
    cover_docx.save(cover_path)

    body_path = tmp_path / "body.docx"
    body_docx = Document()
    body_docx.add_paragraph("DOCUMENTO FINAL")
    body_docx.save(body_path)

    monkeypatch.setattr(
        docx_service,
        "urlopen",
        lambda *_args, **_kwargs: FakeUrlResponse(cover_path.read_bytes()),
    )

    service = docx_service.DocxService.__new__(docx_service.DocxService)
    service.backend_url = "http://backend.test"
    service._prepend_uploaded_cover_docx(
        body_path,
        {"cover_docx_url": "http://backend.test/storage/caratula.docx"},
    )

    paragraphs = [paragraph.text for paragraph in Document(body_path).paragraphs]
    non_empty = [text for text in paragraphs if text]

    assert non_empty[0] == "CARATULA SUBIDA"
    assert "DOCUMENTO FINAL" in non_empty
    assert non_empty.index("CARATULA SUBIDA") < non_empty.index("DOCUMENTO FINAL")


def test_docx_service_resolves_cover_storage_path_when_url_is_missing(tmp_path, monkeypatch) -> None:
    cover_path = tmp_path / "cover.docx"
    cover_docx = Document()
    cover_docx.add_paragraph("CARATULA DESDE STORAGE")
    cover_docx.save(cover_path)

    body_path = tmp_path / "body.docx"
    body_docx = Document()
    body_docx.add_paragraph("DOCUMENTO FINAL")
    body_docx.save(body_path)

    requested_urls = []

    def fake_urlopen(url, *args, **kwargs):
        requested_urls.append(url)
        return FakeUrlResponse(cover_path.read_bytes())

    monkeypatch.setattr(docx_service, "urlopen", fake_urlopen)

    service = docx_service.DocxService.__new__(docx_service.DocxService)
    service.backend_url = "http://backend.test"
    service._prepend_uploaded_cover_docx(
        body_path,
        {"cover_docx_storage_path": "tesis/tesis-1/caratula/caratula demo.docx"},
    )

    non_empty = [paragraph.text for paragraph in Document(body_path).paragraphs if paragraph.text]

    assert requested_urls == [
        "http://backend.test/storage/tesis/tesis-1/caratula/caratula%20demo.docx"
    ]
    assert non_empty[0] == "CARATULA DESDE STORAGE"


def test_docx_service_prepends_uploaded_pdf_cover(tmp_path, monkeypatch) -> None:
    import fitz

    cover_path = tmp_path / "cover.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "CARATULA PDF")
    pdf.save(cover_path)
    pdf.close()

    body_path = tmp_path / "body.docx"
    body_docx = Document()
    body_docx.add_paragraph("DOCUMENTO FINAL")
    body_docx.save(body_path)

    monkeypatch.setattr(
        docx_service,
        "urlopen",
        lambda *_args, **_kwargs: FakeUrlResponse(cover_path.read_bytes()),
    )

    service = docx_service.DocxService.__new__(docx_service.DocxService)
    service.backend_url = "http://backend.test"
    service._prepend_uploaded_cover_docx(
        body_path,
        {
            "cover_docx_url": "http://backend.test/storage/caratula.pdf",
            "cover_docx_original_name": "caratula.pdf",
            "cover_docx_mime_type": "application/pdf",
        },
    )

    document = Document(body_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    with ZipFile(body_path) as archive:
        document_xml = archive.read("word/document.xml").decode()

    assert len(document.inline_shapes) >= 1
    assert "DOCUMENTO FINAL" in document_text
    assert document_xml.index("<w:drawing>") < document_xml.index("DOCUMENTO FINAL")


def test_reference_routes_use_repository(monkeypatch) -> None:
    monkeypatch.setattr(references, "get_repository", lambda: FakeReferencesRepository())
    tesis_id = uuid4()
    reference_id = uuid4()

    async def scenario():
        async with make_client() as client:
            created_response = await client.post(
                f"/theses/{tesis_id}/references",
                json={
                    "authors": [{"first_name": "Ada", "last_name": "Lovelace"}],
                    "year": 1843,
                    "title": "Notas",
                    "type": "book",
                    "publisher": "Demo",
                },
            )
            deleted_response = await client.delete(f"/references/{reference_id}")
            updated_response = await client.patch(
                f"/references/{reference_id}",
                json={"title": "Notas actualizadas"},
            )
        return created_response.json(), updated_response.json(), deleted_response.json()

    created, updated, deleted = run(scenario())

    assert created["title"] == "Notas"
    assert updated["title"] == "Notas actualizadas"
    assert deleted == {"id": str(reference_id), "deleted": True}
