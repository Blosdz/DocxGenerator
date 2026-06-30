from datetime import UTC, datetime
from uuid import uuid4
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

from app.models.references import Author, ReferenceRead, ReferenceType
from app.models.thesis import SectionRead, SubsectionBase, ThesisRead
from app.services.docx_service import DocxService
from app.services.template_service import TemplateService


def test_template_service_recreates_empty_apa7_template(tmp_path) -> None:
    template_path = tmp_path / "apa7.docx"
    template_path.write_bytes(b"")
    service = TemplateService()
    service.templates_dir = tmp_path

    resolved = service.get_template_path()

    assert resolved == template_path
    assert template_path.stat().st_size > 0
    Document(template_path)


def test_docx_generation_writes_heading_levels(tmp_path) -> None:
    template_service = TemplateService()
    template_service.templates_dir = tmp_path / "templates"
    docx_service = DocxService(template_service=template_service)
    docx_service.generated_dir = tmp_path / "generated"
    tesis_id = uuid4()
    now = datetime.now(UTC)

    thesis = ThesisRead(
        id=tesis_id,
        data={
            "metadata": {
                "title": "Tesis demo",
                "author": "Ada Lovelace",
                "institution": "Universidad Demo",
            }
        },
        version=1,
        created_at=now,
        updated_at=now,
    )
    sections = [
        SectionRead(
            id=uuid4(),
            tesis_id=tesis_id,
            title="Capítulo I",
            subtitle=None,
            level=1,
            content="Contenido inicial.",
            order=1,
            version=1,
            created_at=now,
            updated_at=now,
        ),
        SectionRead(
            id=uuid4(),
            tesis_id=tesis_id,
            title="Problema general",
            subtitle=None,
            level=2,
            content="Detalle.",
            order=2,
            version=1,
            created_at=now,
            updated_at=now,
        ),
        SectionRead(
            id=uuid4(),
            tesis_id=tesis_id,
            title="Problema específico",
            subtitle=None,
            level=3,
            content="Detalle específico.",
            order=3,
            version=1,
            created_at=now,
            updated_at=now,
        ),
    ]
    reference_id = uuid4()
    references = [
        ReferenceRead(
            id=reference_id,
            tesis_id=tesis_id,
            authors=[Author(first_name="Roberto", last_name="Hernández")],
            year=2014,
            title="Metodología de la investigación",
            type=ReferenceType.BOOK,
            publisher="McGraw-Hill",
            version=1,
            created_at=now,
            updated_at=now,
        )
    ]
    sections[0].content = f"Contenido inicial con cita [cite:{reference_id}]."
    sections[0].subsections = [
        SubsectionBase(title="Antecedentes", content="Contenido del subheading."),
        SubsectionBase(
            title=f"Marco teórico [cite:{reference_id}]",
            content=f"Contenido con otra cita {{{{cite:{reference_id}}}}}.",
        ),
    ]

    response = docx_service.generate(thesis, sections, references)
    document = Document(response.path)
    styles_by_text = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}
    heading_style = document.styles["Heading 1"]
    normal_style = document.styles["Normal"]

    assert response.path.exists()
    assert document.settings.element.find(qn("w:updateFields")).get(qn("w:val")) == "true"
    assert document.core_properties.title == "Tesis demo"
    assert normal_style.font.name == "Times New Roman"
    assert normal_style.paragraph_format.space_before.pt == 0
    assert normal_style.paragraph_format.space_after.pt == 0
    assert heading_style.font.name == "Times New Roman"
    assert heading_style.font.bold is True
    assert heading_style.font.size.pt == 12
    assert styles_by_text["Tesis demo"] == "Title"
    assert styles_by_text["Tabla de contenido"] != "Heading 1"
    assert styles_by_text["1. Capítulo I"] == "Heading 1"
    assert styles_by_text["1.1 Antecedentes"] == "Heading 2"
    assert styles_by_text["2. Problema general"] == "Heading 2"
    assert styles_by_text["3. Problema específico"] == "Heading 3"
    assert styles_by_text["Contenido del subheading."] == "Normal"

    with ZipFile(response.path) as archive:
        document_xml = archive.read("word/document.xml").decode()
        sources_xml = archive.read("customXml/item1.xml")
        sources_xml_text = sources_xml.decode()
        document_rels_xml = archive.read("word/_rels/document.xml.rels").decode()
        item_rels_xml = archive.read("customXml/_rels/item1.xml.rels").decode()

    source_root = ET.fromstring(sources_xml)
    namespace = {"b": "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"}
    source = source_root.find("b:Source", namespace)

    assert "TOC" in document_xml
    assert '\\o "1-3"' in document_xml
    assert "\\h" in document_xml
    assert "Right-click to update field." not in document_xml
    assert document_xml.count("1. Capítulo I") >= 2
    assert document_xml.count("1.1 Antecedentes") >= 2
    assert "BIBLIOGRAPHY" in document_xml
    assert "1.2 Marco teórico" in document_xml
    assert "<w:fldSimple" in document_xml
    assert (
        f'w:instr="CITATION Ref_{reference_id.hex[:16]} \\m Ref_{reference_id.hex[:16]}"'
        in document_xml
    )
    assert document_xml.count("<w:fldSimple") >= 3
    assert "(Hernández, 2014)" in document_xml
    assert "Actualice el campo en Word para generar la bibliografía APA." not in document_xml
    assert "Hernández, R. (2014). Metodología de la investigación. McGraw-Hill." in document_xml
    assert source_root.get("SelectedStyle") == "/APA.XSL"
    assert source is not None
    assert source.findtext("b:SourceType", namespaces=namespace) == "Book"
    assert source.findtext("b:Tag", namespaces=namespace) == f"Ref_{reference_id.hex[:16]}"
    assert source.findtext("b:Title", namespaces=namespace) == "Metodología de la investigación"
    assert source.findtext("b:Publisher", namespaces=namespace) == "McGraw-Hill"
    assert 'xmlns="http://schemas.openxmlformats.org/officeDocument/2006/bibliography"' in sources_xml_text
    assert 'Version="6"' in sources_xml_text
    assert "relationships/customXml" in document_rels_xml
    assert "../customXml/item1.xml" in document_rels_xml
    assert "relationships/customXmlProps" in item_rels_xml
    assert "itemProps1.xml" in item_rels_xml


def test_docx_generation_prepends_uploaded_cover_docx(tmp_path, monkeypatch) -> None:
    template_service = TemplateService()
    template_service.templates_dir = tmp_path / "templates"
    docx_service = DocxService(template_service=template_service)
    docx_service.generated_dir = tmp_path / "generated"
    docx_service.enable_docm_macro = False
    tesis_id = uuid4()
    now = datetime.now(UTC)

    cover_path = tmp_path / "cover.docx"
    cover_document = Document()
    cover_document.add_paragraph("CARATULA PERSONALIZADA")
    cover_document.save(cover_path)

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def read(self) -> bytes:
            return cover_path.read_bytes()

    def fake_urlopen(url: str, timeout: int):
        assert url == "http://backend.test/storage/caratula.docx"
        assert timeout == 10
        return FakeResponse()

    monkeypatch.setattr("app.services.docx_service.urlopen", fake_urlopen)

    thesis = ThesisRead(
        id=tesis_id,
        data={
            "metadata": {
                "title": "Tesis con carátula",
                "cover_docx_url": "http://backend.test/storage/caratula.docx",
            }
        },
        version=1,
        created_at=now,
        updated_at=now,
    )

    response = docx_service.generate(thesis, [], [])
    document = Document(response.path)
    non_empty_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    assert non_empty_text[0] == "CARATULA PERSONALIZADA"
    assert "Tabla de contenido" in non_empty_text
    assert "Tesis con carátula" not in non_empty_text[:2]


def test_docx_generation_packages_docm_macro_when_vba_asset_exists(tmp_path) -> None:
    template_service = TemplateService()
    template_service.templates_dir = tmp_path / "templates"
    docx_service = DocxService(template_service=template_service)
    docx_service.generated_dir = tmp_path / "generated"
    docx_service.vba_project_path = tmp_path / "vbaProject.bin"
    docx_service.vba_project_path.write_bytes(b"fake-vba-project-for-package-test")
    tesis_id = uuid4()
    reference_id = uuid4()
    now = datetime.now(UTC)

    thesis = ThesisRead(
        id=tesis_id,
        data={"metadata": {"title": "Tesis con macro", "author": "Ada Lovelace"}},
        version=1,
        created_at=now,
        updated_at=now,
    )
    section = SectionRead(
        id=uuid4(),
        tesis_id=tesis_id,
        title=f"Capítulo con cita {{{{cite:{reference_id}}}}}",
        subtitle=None,
        level=1,
        content=f"Cuerpo con cita [cite:{reference_id}].",
        order=1,
        version=1,
        created_at=now,
        updated_at=now,
    )
    section.subsections = [
        SubsectionBase(
            title=f"Subtítulo [cite:{reference_id}]",
            content=f"Contenido del subtítulo {{{{cite:{reference_id}}}}}.",
        )
    ]
    sections = [section]
    references = [
        ReferenceRead(
            id=reference_id,
            tesis_id=tesis_id,
            authors=[Author(first_name="Roberto", last_name="Hernández")],
            year=2014,
            title="Metodología de la investigación",
            type=ReferenceType.BOOK,
            publisher="McGraw-Hill",
            version=1,
            created_at=now,
            updated_at=now,
        )
    ]

    response = docx_service.generate(thesis, sections, references)

    assert response.filename.endswith(".docm")
    assert response.format == "docm"
    assert response.mime_type == "application/vnd.ms-word.document.macroEnabled.12"

    with ZipFile(response.path) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode()
        content_types_xml = archive.read("[Content_Types].xml").decode()
        document_rels_xml = archive.read("word/_rels/document.xml.rels").decode()
        sources_xml = archive.read("customXml/item1.xml").decode()

    reference_tag = f"Ref_{reference_id.hex[:16]}"
    assert "word/vbaProject.bin" in names
    assert "application/vnd.ms-word.document.macroEnabled.main+xml" in content_types_xml
    assert "application/vnd.ms-office.vbaProject" in content_types_xml
    assert "relationships/vbaProject" in document_rels_xml
    assert "vbaProject.bin" in document_rels_xml
    assert "relationships/customXml" in document_rels_xml
    assert "../customXml/item1.xml" in document_rels_xml
    assert "BIBLIOGRAPHY" in document_xml
    assert "TOC" in document_xml
    assert "Right-click to update field." not in document_xml
    assert document_xml.count("<w:fldSimple") >= 4
    assert f'w:instr="CITATION {reference_tag} \\m {reference_tag}"' in document_xml
    assert "(Hernández, 2014)" in document_xml
    assert f"<b:Tag>{reference_tag}</b:Tag>" in sources_xml
